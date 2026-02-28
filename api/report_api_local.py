import os
import datetime
from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import requests
app = Flask(__name__)

# ===============================
# 🔹 DeepSeek 报告提示词构建
# ===============================
# 将 DeepSeek 输出 JSON 转换为 report_api 可用格式
def parse_deepseek_json(deepseek_json, px_to_m, start_depth_m):
    """
    解析 DeepSeek 智能体输出 JSON：
    - 裂缝：仅来自 U-Net
    - 孔洞：来自 vug_results
    - 深度统一由 px → m 换算
    """

    # ================= 裂缝解析（仅 U-Net） =================
    fractures = []

    for r in deepseek_json.get("unet_results", []):
        for m in r.get("metrics_list", []):
            y_offset = m.get("y_offset", 0)
            D_px = m.get("D", 0)

            depth_m = start_depth_m + (y_offset + D_px) * px_to_m

            fractures.append({
                "length_mm": m.get("length_mm", 0),
                "dip_angle_deg": m.get("倾角_deg", 0),
                "depth_m": round(depth_m, 3),
                "area_mm2": m.get("area_mm2", 0),
                "source": "unet"
            })

        # ================= 孔洞解析（正确版本） =================
        vugs = []
        vug_results = deepseek_json.get("vug_results")
        if vug_results:
            for v in vug_results.get("window_metrics", []):
                depth_start_mm = v.get("depth_start_mm", None)
                depth_end_mm = v.get("depth_end_mm", None)

                # 防御式校验
                if depth_start_mm is None or depth_end_mm is None:
                    continue

                # 使用窗口中点作为代表深度（工程上最常用）
                depth_m = (depth_start_mm + depth_end_mm) / 2.0 / 1000.0  # mm → m

                vugs.append({
                    "vug_count": v.get("vug_count", 0),
                    "area_mm2": v.get("total_area_mm2", 0),
                    "depth_m": round(depth_m, 3),
                    "CVPA": v.get("CVPA", 0),
                    "CDENS": v.get("CDENS", 0),
                    "CSIZE": v.get("CSIZE", 0),
                    "source": "unet_vug"
                })

    # ================= 汇总报告 =================
    report_json = {
        "timestamp": deepseek_json.get("timestamp"),
        "modules_used": ["YOLO", "UNet", "SAM2"],
        "params_used": deepseek_json.get("params_used", {}),
        "fractures": fractures,
        "vugs": vugs,
        "reliability_score": "高"
    }

    return report_json



def build_prompt(result_json, image_url=None):
    """构造 DeepSeek 综合分析提示词"""
    system_prompt = """
你是一名资深地质工程师，专长于电成像测井解释与裂缝分析。
请根据输入 JSON 数据生成一份结构化、可直接写入报告的“电成像综合地质分析报告”。
这口井是选取的977.5-980深度段，是安山岩
生成报告时请严格遵循以下要求：
1. 使用中文撰写；
2. 按以下结构输出（以 Markdown 格式）：
   - 图件组成与数据说明
   - 储层电成像特征
   - 裂缝识别与发育特征
   - 裂缝与孔洞综合分析
   - 可靠性与模型一致性
   - 地质意义与工程建议
   - 附录（含参数与统计表说明）
3. 对 YOLO、SAM2、UNet、VUG 模块的检测结果进行综合分析；
4. 若 JSON 中包含裂缝与孔洞信息，请分析其深度分布、发育程度及地质意义；
5. 若 JSON 中缺少伽马或岩性信息，请在报告中说明；
6. 提出针对钻完井、压裂或储层预测的定性建议；
7. 输出时保持 Markdown 格式，标题使用“#、##”层级，数值引用 JSON 中的检测结果。
"""

    user_prompt = f"以下是系统检测输出的 JSON 结果：\n\n```json\n{result_json}\n```\n"
    user_prompt += f"请据此生成正式的综合分析报告（附图：{image_url or '无'}）。"
    return system_prompt, user_prompt


# ===============================
# 🔹 DeepSeek 调用与报告生成
# ===============================
def generate_comprehensive_report(result_json, image_path, image_url=None,
                                  output_path="DeepSeek_ComprehensiveReport.docx"):
    """生成 DeepSeek 综合地质分析报告（附详细裂缝/孔洞表格 + 条形图高亮）"""
    #client = OpenAI(
        #api_key="",
        #base_url="")
    system_prompt, user_prompt = build_prompt(result_json, image_url=image_url)

    #response = client.chat.completions.create(
        #model="deepseek-chat",
        #messages=[
            #{"role": "system", "content": system_prompt},
            #{"role": "user", "content": user_prompt}
        #],
        #temperature=0.4,
        #max_tokens=3500
    #)
    # ===============================
    # 🔹 改为本地 Ollama 调用
    # ===============================
    ollama_url = "http://localhost:11434/api/chat"

    payload = {
        "model": "deepseek-r1:14b",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        "options": {
            "temperature": 0.4,
            "num_predict": 3500
        },
        "stream": False
    }

    response = requests.post(ollama_url, json=payload)

    if response.status_code != 200:
        raise Exception(f"Ollama 调用失败: {response.text}")

    report_text = response.json()["message"]["content"].strip()

    #report_text = response.choices[0].message.content.strip()

    # === 写入 Word 报告 ===
    doc = Document()
    doc.add_heading("电成像综合地质分析报告", level=1)
    doc.add_paragraph("（由 DeepSeek 智能模型自动生成）\n")

    # 主体内容（保留 Markdown 风格）
    for line in report_text.split("\n"):
        if line.startswith("#"):
            doc.add_heading(line.replace("#", "").strip(), level=line.count("#"))
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line.strip())

    # === 附图 ===
    if os.path.exists(image_path):
        doc.add_heading("附图：检测结果综合展示", level=2)
        doc.add_picture(image_path, width=Inches(5.5))
    else:
        doc.add_paragraph(f"⚠️ 未找到图像文件：{image_path}")

    # === 附录参数表格 ===
    doc.add_heading("附录：分析元数据与检测统计", level=2)
    params = result_json.get("params_used", {})

    # 基础元数据表
    table_meta = doc.add_table(rows=1, cols=2)
    table_meta.style = "Light List"
    hdr_cells = table_meta.rows[0].cells
    hdr_cells[0].text = "参数"
    hdr_cells[1].text = "数值 / 信息"

    metadata_items = [
        ("图像深度范围 (mm)", params.get("image_height_mm", "未知")),
        ("井径宽度 (mm)", params.get("image_width_mm", "未知")),
        ("分析时间", result_json.get("timestamp", "未知")),
        ("检测模块", ", ".join(result_json.get("modules_used", []))),
        ("检测可靠性", result_json.get("reliability_score", "未知")),
    ]

    for key, value in metadata_items:
        row_cells = table_meta.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    # === 裂缝详细统计表（带条形图可视化长度） ===
    fractures = result_json.get("fractures", [])
    if fractures:
        doc.add_heading("裂缝详细统计", level=2)
        table_f = doc.add_table(rows=1, cols=6)
        table_f.style = "Light List"
        hdr_cells = table_f.rows[0].cells
        hdr_cells[0].text = "编号"
        hdr_cells[1].text = "长度 (mm)"
        hdr_cells[2].text = "倾角 (°)"
        hdr_cells[3].text = "位置深度 (mm)"
        hdr_cells[4].text = "面积 (mm²)"
        hdr_cells[5].text = "长度可视化"

        # 获取最大长度，用于条形比例
        max_length = max(f.get("length", 1) for f in fractures)

        for idx, f in enumerate(fractures, 1):
            row_cells = table_f.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(f.get("length", "未知"))
            row_cells[2].text = str(f.get("dip_angle", "未知"))
            row_cells[3].text = str(f.get("depth", "未知"))
            row_cells[4].text = str(f.get("area", "未知"))

            # 条形图表示长度，使用 “█” 字符
            length_val = f.get("length", 0)
            bar_count = int((length_val / max_length) * 20) if max_length > 0 else 0
            row_cells[5].text = "█" * bar_count

    # === 孔洞详细统计表（带条形图可视化面积） ===
    vugs = result_json.get("vugs", [])
    if vugs:
        doc.add_heading("孔洞详细统计", level=2)
        table_v = doc.add_table(rows=1, cols=5)
        table_v.style = "Light List"
        hdr_cells = table_v.rows[0].cells
        hdr_cells[0].text = "编号"
        hdr_cells[1].text = "直径 (mm)"
        hdr_cells[2].text = "位置深度 (mm)"
        hdr_cells[3].text = "面积 (mm²)"
        hdr_cells[4].text = "面积可视化"

        max_area = max(v.get("area", 1) for v in vugs)
        for idx, v in enumerate(vugs, 1):
            row_cells = table_v.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(v.get("diameter", "未知"))
            row_cells[2].text = str(v.get("depth", "未知"))
            row_cells[3].text = str(v.get("area", "未知"))

            # 条形图表示面积
            area_val = v.get("area", 0)
            bar_count = int((area_val / max_area) * 20) if max_area > 0 else 0
            row_cells[4].text = "█" * bar_count

    # 设置表格样式（居中 + 中文字体）
    for table in [table_meta, table_f if fractures else None, table_v if vugs else None]:
        if not table:
            continue
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    paragraph.style.font.name = '微软雅黑'
                    paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 保存文件
    doc.save(output_path)
    return report_text, output_path


# ===============================
# 🔹 Flask API 接口
# ===============================
@app.route("/generate_comprehensive_report", methods=["POST"])
def generate_comprehensive_report_api():
    """输入全井检测 JSON 结果与图像路径，输出 DeepSeek 综合报告"""
    data = request.json
    result_json = data.get("result", {})
    image_path = data.get("image_path", "")
    image_url = data.get("image_url", "")

    report_text, report_path = generate_comprehensive_report(result_json, image_path, image_url)

    return jsonify({
        "status": "success",
        "message": "综合报告已生成",
        "report_path": report_path,
        "report_preview": report_text[:800],
        "timestamp": datetime.datetime.now().isoformat()
    })


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=9095)
