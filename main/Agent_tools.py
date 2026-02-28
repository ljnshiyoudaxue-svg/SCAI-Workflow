from matplotlib import pyplot as plt
from matplotlib.ticker import FuncFormatter
from api_caller import call_sam2_api
import base64
from PIL import Image, ImageDraw
import matplotlib
matplotlib.use("Agg")
import math
from openai import OpenAI
import os
import cv2
import json
import numpy as np
import threading
import os
import requests
from typing import Dict, Any, Optional
from typing import Dict, Any, List

import json
import os
import datetime
import uuid
#一、DeepSeek / LLM 决策与复核相关
# 将 DeepSeek 输出 JSON 转换为 report_api 可用格式
import json
import datetime


def _safe_load(obj, name=""):
    """
    工程级兜底：
    - dict / list → 原样返回
    - str(json) → json.loads
    - 其他 → None
    """
    if isinstance(obj, (dict, list)):
        return obj
    if isinstance(obj, str):
        try:
            return json.loads(obj)
        except Exception:
            print(f"⚠️ {_safe_load.__name__}: failed to load {name}")
            return None
    return None


def parse_deepseek_json(deepseek_json, px_to_m=0.001, start_depth_m=4000):
    """
    解析 DeepSeek 智能体输出 JSON（Executor 退化安全版）

    - 裂缝：来自 UNet（fracture）
    - 孔洞：来自 vug_results
    - 深度：px → m 统一换算
    """

    # ======================================================
    # 0️⃣ 顶层 JSON 修复
    # ======================================================
    print("ENTRY deepseek_json type:", type(deepseek_json))

    deepseek_json = _safe_load(deepseek_json, "deepseek_json")
    if deepseek_json is None:
        raise ValueError("❌ deepseek_json is invalid")

    # Executor 常见包一层 report.result_json
    if "deepseek_json" in deepseek_json:
        deepseek_json = deepseek_json["deepseek_json"]

    deepseek_json = _safe_load(deepseek_json, "deepseek_json")
    print("deepseek_json final type:", type(deepseek_json))

    # ======================================================
    # 1️⃣ YOLO（仅检查结构，不参与定量）
    # ======================================================
    yolo = _safe_load(deepseek_json.get("yolo_result", {}), "yolo_result")
    print("yolo_result type:", type(yolo))

    detections = []
    if isinstance(yolo, dict):
        detections = _safe_load(yolo.get("detections", []), "detections") or []

    print("detections type:", type(detections))
    if isinstance(detections, list):
        for i, d in enumerate(detections[:3]):
            print(f"detections[{i}] type:", type(d))

    # ======================================================
    # 2️⃣ 裂缝解析（UNet）
    # ======================================================
    fractures = []

    unet_results = deepseek_json.get("unet_results")
    unet_results = _safe_load(unet_results, "unet_results")

    print("unet_results type:", type(unet_results))

    all_metrics = []

    def collect_metrics(obj):
        """
        递归收集所有 metrics_list
        """
        if isinstance(obj, dict):
            for k, v in obj.items():
                if k == "metrics_list" and isinstance(v, list):
                    all_metrics.extend(v)
                else:
                    collect_metrics(v)
        elif isinstance(obj, list):
            for item in obj:
                collect_metrics(item)

    collect_metrics(unet_results)

    print(f"🧬 collected metrics count = {len(all_metrics)}")

    for m in all_metrics:
        if not isinstance(m, dict):
            continue

        y_offset = m.get("y_offset", 0)
        D_px = m.get("D", 0)

        depth_m = start_depth_m + (y_offset + D_px) * px_to_m

        fractures.append({
            "length_mm": m.get("length_mm", 0),
            "dip_angle_deg": m.get("倾角_deg", 0),
            "depth_m": round(depth_m, 3),
            "area_mm2": m.get("area_mm2", 0),
            "strike_deg": m.get("走向_deg"),
            "source": m.get("_model_id", "unet")
        })

    print(f"🧱 Parsed fractures: {len(fractures)}")

    # ======================================================
    # 3️⃣ 孔洞解析（VUG）
    # ======================================================
    vugs = []

    vug_results = deepseek_json.get("vug_results")
    vug_results = _safe_load(vug_results, "vug_results")

    print("vug_results type:", type(vug_results))

    if isinstance(vug_results, list):
        for model_vug in vug_results:
            if not isinstance(model_vug, dict):
                continue
            window_metrics = _safe_load(model_vug.get("window_metrics", []), "vug.window_metrics") or []

            for v in window_metrics:
                if not isinstance(v, dict):
                    continue

                depth_start_mm = v.get("depth_start_mm")
                depth_end_mm = v.get("depth_end_mm")
                if depth_start_mm is None or depth_end_mm is None:
                    continue

                depth_m = (depth_start_mm + depth_end_mm) / 2.0 / 1000.0

                vugs.append({
                    "vug_count": v.get("vug_count", 0),
                    "area_mm2": v.get("total_area_mm2", 0),
                    "depth_m": round(depth_m, 3),
                    "CVPA": v.get("CVPA", 0),
                    "CDENS": v.get("CDENS", 0),
                    "CSIZE": v.get("CSIZE", 0),
                    "source": model_vug.get("model_id", "unet_vug")
                })
    elif isinstance(vug_results, dict):
        # 保留对旧格式的兼容
        window_metrics = _safe_load(vug_results.get("window_metrics", []), "vug.window_metrics") or []
        for v in window_metrics:
            if not isinstance(v, dict):
                continue
            depth_start_mm = v.get("depth_start_mm")
            depth_end_mm = v.get("depth_end_mm")
            if depth_start_mm is None or depth_end_mm is None:
                continue
            depth_m = (depth_start_mm + depth_end_mm) / 2.0 / 1000.0
            vugs.append({
                "vug_count": v.get("vug_count", 0),
                "area_mm2": v.get("total_area_mm2", 0),
                "depth_m": round(depth_m, 3),
                "CVPA": v.get("CVPA", 0),
                "CDENS": v.get("CDENS", 0),
                "CSIZE": v.get("CSIZE", 0),
                "source": "unet_vug"
            })

    print(f"🕳️ Parsed vugs: {len(vugs)}")

    # ======================================================
    # 4️⃣ 汇总输出
    # ======================================================
    result = {
        "timestamp": deepseek_json.get(
            "timestamp", datetime.datetime.now().isoformat()
        ),
        "modules_used": ["YOLO", "UNet", "SAM2"],
        "params_used": deepseek_json.get("params_used", {}),
        "fractures": fractures,
        "vugs": vugs,
        "reliability_score": "高"
    }

    return result

def generate_comprehensive_report2(
    result: Any,
    final_image: str,
    log=print,
    report_api_url: str = "http://127.0.0.1:9095/generate_comprehensive_report",
    timeout: int = 120
) -> Dict[str, Optional[str]]:
    """
    Generate comprehensive geological report via report service.
    Returns structured output compatible with Executor PLAN:
    {
        "report": {
            "report_path": str or None,
            "report_preview": str or None
        },
        "status": "success" | "fallback" | "failed"
    }
    """

    # =========================================================
    # 0️⃣ logging 兼容
    # =========================================================
    if isinstance(log, str):
        if log == "print":
            log = print
        else:
            raise ValueError(f"Unknown log function: {log}")

    # =========================================================
    # 1️⃣ 输入检查 & 标准化（🔥 关键）
    # =========================================================
    if not isinstance(result, dict):
        raise ValueError("result must be dict")

    # --- params_used ---
    params_used = result.get("params_used")
    if not isinstance(params_used, dict):
        params_used = {}

    # --- fractures ---
    fractures = result.get("fractures") or []
    if not isinstance(fractures, list):
        fractures = []

    # 给 fracture 补“报告友好字段”
    normalized_fractures = []
    for f in fractures:
        if not isinstance(f, dict):
            continue
        normalized_fractures.append({
            "length_mm": f.get("length_mm", 0),
            "area_mm2": f.get("area_mm2", 0),
            "depth_m": f.get("depth_m"),
            "dip_angle_deg": f.get("dip_angle_deg"),
            "dip": f.get("dip_angle_deg"),          # ⭐ 别名
            "strike_deg": f.get("strike_deg"),
            "azimuth": f.get("strike_deg"),         # ⭐ 别名
            "source": f.get("source", "unknown")
        })

    # --- vugs ---
    vugs = result.get("vugs") or []
    print(vugs)
    if not isinstance(vugs, list):
        vugs = []

    # --- modules_used ---
    modules_used = result.get("modules_used") or []
    if not isinstance(modules_used, list):
        modules_used = []

    # =========================================================
    # 2️⃣ 组装“报告安全 payload”
    # =========================================================
    safe_result = {
        "timestamp": result.get("timestamp"),
        "modules_used": modules_used,
        "params_used": params_used,
        "fractures": normalized_fractures,
        "vugs": vugs,
        "reliability_score": result.get("reliability_score", "未知")
    }

    # debug
    log("📦 report.result_json keys:", safe_result.keys())
    log("fractures count:", len(safe_result["fractures"]))
    log("vugs count:", len(safe_result["vugs"]))

    payload = {
        "result": safe_result,
        "image_path": final_image,
        "image_url": final_image
    }

    # =========================================================
    # 3️⃣ 调用报告服务
    # =========================================================
    try:
        log("🔹 调用报告生成接口 ...")
        import requests
        import os

        report_resp = requests.post(
            report_api_url,
            json=payload,
            timeout=timeout
        )

        if report_resp.status_code != 200:
            log(f"❌ 报告生成失败: HTTP {report_resp.status_code}")
            return {
                "report": {"report_path": None, "report_preview": None},
                "status": "failed"
            }

        # =====================================================
        # 4️⃣ JSON 响应
        # =====================================================
        try:
            resp_json = report_resp.json()
            report_out = {
                "report": {
                    "report_path": resp_json.get("report_path"),
                    "report_preview": resp_json.get("report_preview")
                },
                "status": "success"
            }

            # 输出调试信息
            print("📤 report status:", report_out.get("status"))
            print("📄 report path:", report_out["report"]["report_path"])

            return report_out

        # =====================================================
        # 5️⃣ fallback：非 JSON → Word
        # =====================================================
        except Exception:
            log("⚠️ 报告接口返回非 JSON，使用 fallback Word")

            report_text = report_resp.text.strip()

            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading("电成像综合地质分析报告", level=1)

            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("#"):
                    doc.add_heading(
                        line.replace("#", "").strip(),
                        level=min(line.count("#"), 4)
                    )
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            if final_image and os.path.exists(final_image):
                doc.add_picture(final_image, width=Inches(5.5))

            report_path = "DeepSeek_Report_TCA3.docx"

            doc.save(report_path)

            report_out = {
                "report": {
                    "report_path": report_path,
                    "report_preview": report_text[:500]
                },
                "status": "fallback"
            }

            print("📤 report status:", report_out.get("status"))
            print("📄 report path:", report_out["report"]["report_path"])

            return report_out

    except Exception as e:
        log(f"⚠️ 报告生成时出现异常: {e}")
        return {
            "report": {"report_path": None, "report_preview": None},
            "status": "failed"
        }
def generate_comprehensive_report(
    result: Any,
    final_image: str,
    log=print,
    report_api_url: str = "http://127.0.0.1:9095/generate_comprehensive_report",
    timeout: int = 120
) -> Dict[str, Optional[str]]:
    """
    Generate comprehensive geological report via report service.
    Returns structured output compatible with Executor PLAN:
    {
        "report.report_path": str or None,
        "report.report_preview": str or None,
        "report.status": "success" | "fallback" | "failed"
    }
    """

    # =========================================================
    # 0️⃣ logging 兼容
    # =========================================================
    if isinstance(log, str):
        if log == "print":
            log = print
        else:
            raise ValueError(f"Unknown log function: {log}")

    # =========================================================
    # 1️⃣ 输入检查 & 标准化（🔥 关键）
    # =========================================================
    if not isinstance(result, dict):
        raise ValueError("result must be dict")

    # --- params_used ---
    params_used = result.get("params_used")
    if not isinstance(params_used, dict):
        params_used = {}

    # --- fractures ---
    fractures = result.get("fractures") or []
    if not isinstance(fractures, list):
        fractures = []

    # 给 fracture 补“报告友好字段”
    normalized_fractures = []
    for f in fractures:
        if not isinstance(f, dict):
            continue
        normalized_fractures.append({
            "length_mm": f.get("length_mm", 0),
            "area_mm2": f.get("area_mm2", 0),
            "depth_m": f.get("depth_m"),
            "dip_angle_deg": f.get("dip_angle_deg"),
            "dip": f.get("dip_angle_deg"),          # ⭐ 别名
            "strike_deg": f.get("strike_deg"),
            "azimuth": f.get("strike_deg"),         # ⭐ 别名
            "source": f.get("source", "unknown")
        })

    # --- vugs ---
    vugs = result.get("vugs") or []
    print(vugs)
    if not isinstance(vugs, list):
        vugs = []

    # --- modules_used ---
    modules_used = result.get("modules_used") or []
    if not isinstance(modules_used, list):
        modules_used = []

    # =========================================================
    # 2️⃣ 组装“报告安全 payload”
    # =========================================================
    safe_result = {
        "timestamp": result.get("timestamp"),
        "modules_used": modules_used,
        "params_used": params_used,
        "fractures": normalized_fractures,
        "vugs": vugs,
        "reliability_score": result.get("reliability_score", "未知")
    }

    # debug
    log("📦 report.result_json keys:", safe_result.keys())
    log("fractures count:", len(safe_result["fractures"]))
    log("vugs count:", len(safe_result["vugs"]))

    payload = {
        "result": safe_result,
        "image_path": final_image,
        "image_url": final_image
    }

    # =========================================================
    # 3️⃣ 调用报告服务
    # =========================================================
    try:
        log("🔹 调用报告生成接口 ...")
        import requests
        import os

        report_resp = requests.post(
            report_api_url,
            json=payload,
            timeout=timeout
        )

        if report_resp.status_code != 200:
            log(f"❌ 报告生成失败: HTTP {report_resp.status_code}")
            return {
                "report.report_path": None,
                "report.report_preview": None,
                "report.status": "failed"
            }

        # =====================================================
        # 4️⃣ JSON 响应
        # =====================================================
        try:
            resp_json = report_resp.json()
            report_out = {
                "report.report_path": resp_json.get("report_path"),
                "report.report_preview": resp_json.get("report_preview"),
                "report.status": "success"
            }

            # 输出调试信息
            print("📤 report status:", report_out["report.status"])
            print("📄 report path:", report_out["report.report_path"])

            return report_out

        # =====================================================
        # 5️⃣ fallback：非 JSON → Word
        # =====================================================
        except Exception:
            log("⚠️ 报告接口返回非 JSON，使用 fallback Word")

            report_text = report_resp.text.strip()

            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading("电成像综合地质分析报告", level=1)

            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("#"):
                    doc.add_heading(
                        line.replace("#", "").strip(),
                        level=min(line.count("#"), 4)
                    )
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            if final_image and os.path.exists(final_image):
                doc.add_picture(final_image, width=Inches(5.5))

            report_path = "DeepSeek_Report_TCA3.docx"

            doc.save(report_path)

            report_out = {
                "report.report_path": report_path,
                "report.report_preview": report_text[:500],
                "report.status": "fallback"
            }

            print("📤 report status:", report_out["report.status"])
            print("📄 report path:", report_out["report.report_path"])

            return report_out

    except Exception as e:
        log(f"⚠️ 报告生成时出现异常: {e}")
        return {
            "report.report_path": None,
            "report.report_preview": None,
            "report.status": "failed"
        }
# ===== DeepSeek 决策 =====

import json
import re
from local_llm import call_local_llm
def safe_json_extract(text):
    """
    处理 ```json 包裹 或 多余文本
    """
    text = text.strip()

    # 去掉 ```json ``` 包裹
    if text.startswith("```"):
        text = re.sub(r"```.*?\n", "", text)
        text = text.replace("```", "")

    # 提取第一个 JSON 对象
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        return match.group(0)

    return text

import json

def deepseek_filter_curves_safe(curves_metrics, x_points_list, image_height_px, image_width_px,
                                min_points=200, max_retries=3, flags=None, log_fn=None, strategy="default"):
    """
    ✅ 合并 DeepSeek 曲线复核 + metrics 重写
    - Step 1: 预处理
    - Step 2: 调用本地 LLM (替代原 DeepSeek API)
    - Step 3: 兜底逻辑
    - Step 4: 根据 flags 决定是否覆盖
    """
    if log_fn:
        log_fn("🚀 Start deepseek_filter_and_overwrite")

    # ---------- Step 1: 预处理 ----------
    if not curves_metrics or not x_points_list:
        if log_fn:
            log_fn("⚠️ curves_metrics 或 x_points_list 为空")
        filtered_curves, analysis_log, pre_filtered = [], [], []
    else:
        pre_filtered = [(i, m) for i, m in enumerate(curves_metrics)
                        if i < len(x_points_list) and len(x_points_list[i]) >= min_points]
        if log_fn:
            log_fn(f"🟢 pre_filtered 数量: {len(pre_filtered)}")

        if not pre_filtered:
            filtered_curves, analysis_log = [], ["⚠️ 没有曲线满足 min_points 条件"]
        else:
            # ---------- Step 2: 【修改点】使用本地 Ollama / LLM 替代 API ----------
            metrics_simple = [{"A": m.get("A"), "B": m.get("B"), "C": m.get("C"), "D": m.get("D")} for _, m in pre_filtered]
            try:
                params_json = json.dumps(metrics_simple, ensure_ascii=False)
            except:
                params_json = json.dumps(metrics_simple, ensure_ascii=True)

            prompt = generate_prompt("filter_curves", curves_params=params_json,
                                     image_width_px=image_width_px, image_height_px=image_height_px,
                                     strategy=strategy)

            if log_fn:
                log_fn(f"🔹 Prompt 发送给本地 LLM:\n{prompt[:500]}...")  # 只打印前500字符防止过长

            # 🔹 修改点：调用本地 LLM
            try:
                raw_output = call_local_llm(prompt, temperature=0)  # 【修改点】
                if log_fn:
                    log_fn(f"🔹 LLM raw output:\n{str(raw_output)[:500]}...")  # 前500字符
                cleaned_output = safe_json_extract(raw_output)      # 【修改点】
                if log_fn:
                    log_fn(f"🔹 cleaned_output:\n{cleaned_output[:500]}...")
                decision = json.loads(cleaned_output)              # 【修改点】

                valid_idx = decision.get("valid_curves", [])
                analysis_log = decision.get("analysis_log", [])
                filtered_curves = [pre_filtered[i][1] for i in valid_idx if i < len(pre_filtered)]

                if log_fn:
                    log_fn(f"🟢 LLM 返回有效曲线数量: {len(filtered_curves)}")

            except Exception as e:
                if log_fn:
                    log_fn(f"⚠️ 本地 LLM 尝试失败: {e}")
                    log_fn(f"🔹 raw_output: {repr(raw_output) if 'raw_output' in locals() else 'None'}")
                filtered_curves, analysis_log = [], []

            # ---------- Step 3: 兜底逻辑 ----------
            if not filtered_curves:
                for idx, m in pre_filtered:
                    A, B = abs(m.get("A", 0)), m.get("B", 1.0)
                    valid_mark = "valid" if B <= 0.05 and A <= image_height_px / 3 else "invalid"
                    if valid_mark == "valid":
                        filtered_curves.append(m)
                    analysis_log.append(f"曲线{idx}: B={B}, |A|={A} => {valid_mark}")

    # ---------- Step 4: 根据 flags 决定是否覆盖 ----------
    enable_reflection = False
    if flags and isinstance(flags, dict):
        enable_reflection = flags.get("enable_reflection", False)

    if enable_reflection:
        final_metrics = filtered_curves
        if log_fn:
            log_fn(f"🔄 覆盖 metrics, 数量={len(final_metrics)}")
    else:
        final_metrics = curves_metrics
        if log_fn:
            log_fn(f"⚠️ flags 未启用，保持原 metrics, 数量={len(final_metrics)}")

    # ---------- Step 5: 返回 Executor 格式 ----------
    return {
        "fracture_metrics_list": final_metrics,
        "reflection.curves_filtered": filtered_curves,
        "reflection.analysis_log": analysis_log,
        "reflection.pre_filtered": pre_filtered
    }

# ===== DeepSeek 决策 =====

def deepseek_decide_models(user_input, yolo_results, strategy="default"):
    """
    本地 Ollama 版本
    Semantic-Constrained Model Selection
    1️⃣ 根据用户输入 + YOLO 结果生成 prompt
    2️⃣ 调用本地 LLM (Ollama 或其他) 得到 JSON 决策
    3️⃣ 返回白名单模型 + 可选参数
    """
    prompt = generate_prompt(
        "decide_models",
        user_input=user_input,
        yolo_results=yolo_results,
        strategy=strategy
    )

    try:
        # 调用本地 LLM
        raw_output = call_local_llm(prompt, temperature=0.2)

        # 尝试提取 JSON 内容
        cleaned_output = safe_json_extract(raw_output)
        decision = json.loads(cleaned_output)

        # 允许的模型白名单
        allowed_models = ["unet_Fracture", "unet_Induced_Fracture", "unet_Vug"]
        models = [m for m in decision.get("models", []) if m in allowed_models]

        # 兜底：至少返回一个可执行模型
        if not models:
            models = ["unet_Fracture"]


        # ✅ 返回结果
        return {
            "model_ids": models
        }

    except Exception as e:
        print(f"⚠️ 本地 LLM 模型决策失败: {e}")
        return {
            "model_ids": ["unet_Fracture"]
        }
def generate_prompt(task_type, user_input=None, yolo_results=None, curves_params=None,
                    image_width_px=None, image_height_px=None, strategy="default"):
    """

    生成不同策略的 Prompt
    task_type: "decide_models" 或 "filter_curves"
    strategy: "default", "zero-shot", "few-shot", "hard-constraint"
    """
    if task_type == "decide_models":
        base = "你是一位图像分割智能助手，根据用户输入和YOLO检测结果自动选择最合适的地质分割模型。请严格返回 JSON。"
        if strategy == "default":
            prompt = f"{base}\n用户输入: {user_input}\nYOLO结果: {yolo_results}"
        elif strategy == "zero-shot":
            prompt = f"{base}\n用户输入: {user_input}\nYOLO结果: {yolo_results}\n只返回JSON，不要解释。"
        elif strategy == "few-shot":
            prompt = f"""{base}
示例：
用户输入: 裂缝
YOLO结果: {{}}
输出: {{
  "models": ["unet_Fracture"],
  "parameters": {{}}
}}
用户输入: 孔洞
YOLO结果: {{}}
输出: {{
  "models": ["unet_Vug"],
  "parameters": {{}}
}}
用户输入: 诱导缝
YOLO结果: {{}}
输出: {{
  "models": ["unet_Induced_Fracture"],
  "parameters": {{}}
}}
用户输入: 检测图片
YOLO结果: ["Fracture", "Vug"]
输出: {{"models": ["unet_Fracture", "unet_Vug"], "parameters": {{}} }}

现在用户输入: {user_input}
YOLO结果: {yolo_results}
请根据示例选择最合适的模型。
"""
        elif strategy == "hard-constraint":
            prompt = f"""你是一位图像分割智能助手，负责根据用户的自然语言意图与YOLO检测结果自动选择最合适的地质分割模型。
你的任务：
1. 如果用户输入中出现“裂缝”，仅使用 ["unet_Fracture"]；
2. 如果出现“孔洞”，仅使用 ["unet_Vug"]；
3. 如果出现“诱导缝”或“钻井诱导裂缝”，仅使用 ["unet_Induced_Fracture"]；
4. 如果用户输入为“检测图片”或“分析图片”，则根据 YOLO 检测结果自动匹配：
   - YOLO检测结果中包含 Fracture → 使用 ["unet_Fracture"]
   - YOLO检测结果中包含 Induced_Fracture → 使用 ["unet_Induced_Fracture"]
   - YOLO检测结果中包含 Vug → 使用 ["unet_Vug"]
5. 如果检测结果包含多种类型，可同时输出多个模型；
6. 如果均不符合条件，则返回 ["unet_Fracture"]。
输出格式固定为 JSON：
{{
  "models": ["unet_Fracture", "unet_Vug"],
  "parameters": {{}}
}}
请只输出JSON内容，不要解释。"""
        else:
            raise ValueError(f"未知策略: {strategy}")

    elif task_type == "filter_curves":
        base = "你是地质图像分析助手，根据曲线参数分析每条曲线是否有效。请严格输出 JSON,不要解释、不要添加说明、不要输出多余文字，只输出 JSON 对象。"
        if strategy in ["default", "zero-shot"]:
            prompt = (
                f"{base}\n"
                "你是一名地质曲线分析助手，请基于地质知识和电成像有关的知识判断每条曲线的有效性。\n\n"
                "【输出要求】\n"
                "1. 输出必须是合法的 JSON 对象，不能包含任何解释性文字、注释或额外说明。\n"
                "2. JSON 对象需包含以下三个键：\n"
                "   - \"analysis_log\": 包含每条曲线的分析说明。\n"
                "   - \"valid_curves\": 有效曲线索引列表。\n"
                "   - \"invalid_curves\": 无效曲线索引列表。\n"
                "3. 若无法判断，请返回：\n"
                "{{\"analysis_log\": [], \"valid_curves\": [], \"invalid_curves\": []}}\n\n"
                f"输入数据如下：\n图像宽度={image_width_px}, 高度={image_height_px}, 参数={curves_params}\n\n"
                "请直接输出符合要求的 JSON 对象。"
            )
        elif strategy == "few-shot":
            prompt = (
                f"你是一名地质测井分析专家，请学习以下示例的输入输出格式。\n"
                "任务：判断每条曲线是否有效，并输出 JSON 结果。\n\n"
                "【重要要求】\n"
                "- 你必须只输出 JSON 对象，不能输出任何解释、前缀、注释或自然语言。\n"
                "- 不要输出“输出:”或“结果如下:”，不要添加换行说明。\n"
                "- 输出必须严格遵守 JSON 语法。\n\n"
                "示例：\n"
                "输入: 图像宽度=472, 高度=1422, 参数=[{{\"A\":50,\"B\":0.02}}, {{\"A\":700,\"B\":0.1}},{{\"A\":-66.71,\"B\":0.0132}}, {{\"A\":44,\"B\":0.0132}}, {{\"A\":500,\"B\":0.060}}, {{\"A\":-43.17,\"B\":0.0221}}, {{\"A\":68199.68,\"B\":0.0003 }},{{\"A\":7.62,\"B\":0.0647}} ]\n"
                "输出:\n"
                "{{\n"
                "  \"analysis_log\": [\n"
                "    \"曲线0: A=50, B=0.02 ✅\",\n"
                "    \"曲线1: A=700, B=0.1 ❌\",\n"
                "    \"曲线2: B=0.0132，，A=-66.71, |A|=66.71 ✅\",\n"
                "    \"曲线3: B=0.0132 , A=44, |A|=44  ✅\",\n"
                "    \"曲线4: B=0.060, A=500, |A|=500  ❌ \",\n"
                "    \"曲线5: B=0.0221 , A=-43.17, |A|=43.17 ✅ \",\n"
                "    \"曲线6: B=0.0003 , A=68199.68, |A|=68199.68 ❌ \",\n"
                "    \"曲线7: B=0.0647, A=7.62, |A|=7.62 ❌ \"\n"
                "  ],\n"
                "  \"valid_curves\": [0, 2, 3, 5],\n"
                "  \"invalid_curves\": [1, 4, 6, 7]\n"
                "}}\n\n"
                f"现在输入:\n图像宽度={image_width_px}, 高度={image_height_px}, 参数={curves_params}\n\n"
                "请模仿示例输出，仅返回 JSON 对象。\n"
                "如果无法判断，请返回：\n"
                "{{\"analysis_log\": [], \"valid_curves\": [], \"invalid_curves\": []}}"
            )

        elif strategy == "hard-constraint":
            prompt = f"""
你是地质图像分析助手，负责根据曲线参数判断哪些曲线有效。请严格遵守以下规则生成输出：

【任务要求】
1. 对输入的每条曲线都生成一条日志（即使是无效曲线，也必须生成）。
2. 在处理振幅参数 A 时，请取其绝对值，因为振幅应为正数。
3. 判断标准：
   - 若 B <= 0.02 且 |A| <= 图像高度 * 2/3，则该曲线有效（✅）；
   - 否则视为无效（❌）。

【输出格式】
必须严格返回 JSON 格式，不允许出现除 JSON 外的任何解释或文字。示例如下：
{{
  "analysis_log": ["曲线0: A=100 ✅", "曲线1: A=600 ❌"],
  "valid_curves": [0],
  "invalid_curves": [1]
}}

【输入数据】
图像宽度 = {image_width_px}, 高度 = {image_height_px}, 参数 = {curves_params}

请严格按照上述标准输出 JSON。
"""
        else:
            raise ValueError(f"未知策略: {strategy}")
    else:
        raise ValueError(f"未知任务类型: {task_type}")

    return prompt





#二、YOLO / U-Net / SAM2 模型调用接口
# ===== 调用 U-Net API =====
# ===== U-Net API 配置 =====
# yolo_agent.py
import requests

YOLO_API_URL = "http://localhost:2000/analyze"  # 若部署在远程，改为实际IP:端口

def call_yolo_api(image_path: str, params: dict = {}):
    """
    发送图像到 YOLO Flask API，返回检测结果（包含 class、bbox、confidence）
    :param image_path: 本地图像路径
    :param params: 可选参数，暂未启用（保留扩展）
    :return: dict {detections: [...]}
    """
    with open(image_path, 'rb') as f:
        files = {'image': f}
        try:
            response = requests.post(YOLO_API_URL, files=files, data=params)
            response.raise_for_status()
            return response.json()
        except requests.RequestException as e:
            print(f"❌ 调用YOLO API失败：{e}")
            raise

UNET_API_URL = "http://127.0.0.1:7000/unet/{model_id}/segment"
def call_unet_api(model_id, image_path):
    with open(image_path, "rb") as f:
        files = {"roi": f}
        response = requests.post(UNET_API_URL.format(model_id=model_id), files=files)
    if response.status_code == 200:
        data = response.json()
        mask_base64 = data["mask"]
        mask_bytes = base64.b64decode(mask_base64)
        mask_path = f"{model_id}_{uuid.uuid4().hex}.png"

        #mask_path = os.path.join(temp_dir, f"{model_id}_{uuid.uuid4().hex}.png")
        with open(mask_path, "wb") as f:
            f.write(mask_bytes)
        return {"mask": mask_path}
    else:
        raise RuntimeError(f"❌ U-Net API调用失败: {response.text}")

def call_sam2_box(image_path, box_coords):
    """
    使用 box 提示调用 SAM2 分割
    :param image_path: 输入图像路径
    :param box_coords: [x_min, y_min, x_max, y_max]
    :return: {"mask": mask_path}
    """
    with open(image_path, "rb") as f:
        files = {"file": (image_path, f, "image/jpeg")}
        data = {"prompt_type": "box", "box_coords": str(box_coords)}
        resp = requests.post("http://127.0.0.1:3000/predict", files=files, data=data)
        resp.raise_for_status()
        result = resp.json()

    # 保存 mask 文件
    mask_base64 = result.get("mask")
    mask_bytes = base64.b64decode(mask_base64)
    mask_path = f"sam2_box_{uuid.uuid4().hex}.png"
    #mask_path = os.path.join(temp_dir, f"sam2_box_{uuid.uuid4().hex}.png")
    with open(mask_path, "wb") as f:
        f.write(mask_bytes)

    return {"mask": mask_path}
#三、Mask 与图像预处理
def save_base64_mask(mask_b64, save_path):
    """将 base64 mask 保存为 PNG 文件"""
    with open(save_path, "wb") as f:
        f.write(base64.b64decode(mask_b64))
    return save_path
# ===== Mask 预处理 =====
def preprocess_mask_for_analysis(mask_path, log_fn=None):
    mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
    if mask_img is None:
        raise RuntimeError(f"读取 mask 文件失败: {mask_path}")
    binary_mask = (mask_img > 0).astype(np.uint8) * 255
    nonzero_count = cv2.countNonZero(binary_mask)
    if log_fn:
        log_fn(f"🔹 mask 非零像素数: {nonzero_count}")
        if nonzero_count == 0:
            log_fn("⚠️ 当前 mask 全为空白，分析结果可能无效")
    temp_mask_path = f"temp_mask_{uuid.uuid4().hex}.png"
    #temp_mask_path = os.path.join(temp_dir, f"temp_mask_{uuid.uuid4().hex}.png")
    cv2.imwrite(temp_mask_path, binary_mask)
    return temp_mask_path
def split_mask_to_contours(mask_path):
    mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
    binary_mask = (mask_img > 0).astype(np.uint8)
    contours, _ = cv2.findContours(binary_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    single_masks = []
    for contour in contours:
        mask_single = np.zeros_like(binary_mask)
        cv2.drawContours(mask_single, [contour], -1, 255, -1)
        single_masks.append(mask_single)
    return single_masks
# ===== 颜色映射 =====


# SAM2 专用颜色映射
SAM2_MAPPING = {
    "sam2_fracture": {"color": (0, 255, 0)},   # 紫色
    "sam2_vug": {"color": (255, 165, 0)}         # 橙色
}
MODEL_MAPPING = {
    "unet_Fracture": {"color": (255, 0, 0)},          # 红色
    "unet_Induced_Fracture": {"color": (0, 0, 255)},  # 蓝色
    "unet_Vug": {"color": (0, 255, 0)}               # 绿色
}
def overlay_masks2(masks, model_ids, base_image_path, log_fn=None):
    """
    将 U-Net（裂缝 / 孔洞）+ SAM2 的掩码叠加到原图上，并使用不同颜色区分
    masks: list of dicts or list of strings
    model_ids: list of str
    """
    print(">>> overlay received masks =", masks)
    # ===== ✅ 关键修复：保证 masks 是 list =====
    if isinstance(masks, (str, dict)):
        masks = [masks]

    print(">>> overlay received masks =", masks)
    base_img = cv2.imread(base_image_path)
    if base_img is None:
        raise FileNotFoundError(f"无法读取底图: {base_image_path}")

    overlay = base_img.copy()
    print(model_ids)
    model_ids = model_ids['selected_models']
    print(model_ids)
    for mask, model_id in zip(masks, model_ids):
        model_id_lower = model_id.lower()

        # ---------- 颜色策略 ----------
        if "fracture" in model_id_lower:
            color = (0, 0, 255)        # 🔴 裂缝
        elif "vug" in model_id_lower:
            color = (0, 255, 255)      # 🟡 孔洞
        elif "sam2" in model_id_lower:
            color = (0, 255, 0)        # 🟢 SAM2
        else:
            color = (255, 255, 255)    # ⚪ 兜底

        # ---------- 读取 mask ----------
        if isinstance(mask, dict):
            mask_path = mask.get("mask")
        elif isinstance(mask, str):
            mask_path = mask
        else:
            if log_fn:
                log_fn(f"⚠️ overlay_masks 跳过未知类型: {type(mask)}")
            continue

        if not mask_path or not os.path.exists(mask_path):
            if log_fn:
                log_fn(f"⚠️ mask 不存在: {mask_path}")
            continue

        mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
        print(f"🧩 mask_img.shape = {mask_img.shape}, dtype = {mask_img.dtype}")
        if mask_img is None:
            continue

        # ---------- 二值化 ----------
        _, binary = cv2.threshold(mask_img, 1, 255, cv2.THRESH_BINARY)

        # ---------- 提取轮廓 ----------
        contours, _ = cv2.findContours(
            binary,
            cv2.RETR_EXTERNAL,
            cv2.CHAIN_APPROX_SIMPLE
        )

        # ---------- 绘制（填充） ----------
        cv2.drawContours(
            overlay,
            contours,
            -1,
            color,
            thickness=cv2.FILLED
        )

    out_path = base_image_path.replace(".png", "_overlay.png")
    cv2.imwrite(out_path, overlay)
    print(out_path)
    return out_path
def overlay_masks(masks, model_ids, base_image_path, log_fn=None):
    """
    将 U-Net（裂缝 / 孔洞）+ SAM2 的掩码叠加到原图上，并使用不同颜色区分
    masks: list of dicts or list of strings
    model_ids: list of str 或 dict 包含 selected_models
    """
    print(">>> overlay received masks =", masks)
    import os, cv2

    # ===== ✅ 关键修复：保证 masks 是 list =====
    if isinstance(masks, (str, dict)):
        masks = [masks]

    base_img = cv2.imread(base_image_path)
    if base_img is None:
        raise FileNotFoundError(f"无法读取底图: {base_image_path}")

    overlay = base_img.copy()

    # ===== 处理 model_ids =====
    if isinstance(model_ids, dict) and 'selected_models' in model_ids:
        model_ids = model_ids['selected_models']
    if not isinstance(model_ids, list):
        model_ids = []

    # ===== 遍历 masks，不再用 zip，而是单独匹配颜色 =====
    for mask in masks:
        # ---------- 读取 mask_path ----------
        if isinstance(mask, dict):
            mask_path = mask.get("mask")
        elif isinstance(mask, str):
            mask_path = mask
        else:
            if log_fn:
                log_fn(f"⚠️ overlay_masks 跳过未知类型: {type(mask)}")
            continue

        if not mask_path or not os.path.exists(mask_path):
            if log_fn:
                log_fn(f"⚠️ mask 不存在: {mask_path}")
            continue

        # ---------- 根据文件名或类型匹配颜色 ----------
        color = (255, 255, 255)  # 默认白色
        lower_path = mask_path.lower()
        if "full_mask" in lower_path:
            color = (0, 0, 255)        # 🔴 裂缝
        elif "full_mask_vug" in lower_path:
            color = (0, 255, 255)      # 🟡 孔洞
        else:
            # 尝试用 model_ids 匹配
            for mid in model_ids:
                mid_lower = mid.lower()
                if "fracture" in mid_lower:
                    color = (0, 0, 255)
                    break
                elif "vug" in mid_lower:
                    color = (0, 255, 255)
                    break
                elif "sam2" in mid_lower:
                    color = (0, 255, 0)
                    break

        mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
        if mask_img is None:
            continue

        # ---------- 二值化 ----------
        _, binary = cv2.threshold(mask_img, 1, 255, cv2.THRESH_BINARY)

        # ---------- 提取轮廓 ----------
        contours, _ = cv2.findContours(
            binary,
            cv2.RETR_EXTERNAL,
            cv2.CHAIN_APPROX_SIMPLE
        )

        # ---------- 绘制（填充） ----------
        cv2.drawContours(
            overlay,
            contours,
            -1,
            color,
            thickness=cv2.FILLED
        )

    # ===== 输出结果路径 =====
    out_path = base_image_path.replace(".png", "_overlay.png")
    cv2.imwrite(out_path, overlay)
    if log_fn:
        log_fn(f"✅ Overlay saved: {out_path}")

    return out_path
def visualize_pipeline(
    image_path,
    sliding_results,
    unet_results,
    yolo_results,
    model_ids,
    image_height_px,
    log_fn=None
):
    """
    可视化总调度函数
    顺序：
    1. 提取裂缝/孔洞 mask
    2. 叠加 overlay
    3. 绘制最终检测 + 曲线
    """

    if log_fn:
        log_fn("🚀 Start visualization pipeline")

    # ========= Step 1: 提取 mask =========
    overlay_dict = extract_overlay_masks(
        sliding_results,
        log_fn=log_fn
    )

    masks = [
        overlay_dict.get("derived.overlay.fracture_mask"),
        overlay_dict.get("derived.overlay.vug_mask")
    ]

    # 过滤 None
    masks = [m for m in masks if m]

    if log_fn:
        log_fn(f"🟢 Masks extracted: {masks}")

    # ========= Step 2: 叠加 overlay =========
    overlay_path = None
    if masks:
        overlay_path = overlay_masks(
            masks,
            model_ids,
            image_path,
            log_fn=log_fn
        )
    else:
        overlay_path = image_path
        if log_fn:
            log_fn("⚠️ No masks found, skip overlay")

    # ========= Step 3: 绘制最终结果 =========
    final_path = draw_final_results(
        overlay_path,
        unet_results,
        yolo_results,
        image_height_px
    )

    if log_fn:
        log_fn(f"✅ Visualization done: {final_path}")

    return {
        "overlay_image": overlay_path,
        "final_image": final_path
    }
#四、裂缝 / 孔洞参数分析接口
# ===== 裂缝分析 API =====
def call_crack_api(mask_path, image_height_mm, image_width_mm):
    mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
    binary_mask = (mask_img > 0).astype(np.uint8) * 255
    temp_mask_path = f"temp_mask_{uuid.uuid4().hex}.png"
    #temp_mask_path = os.path.join(temp_dir, f"temp_mask_{uuid.uuid4().hex}.png")
    cv2.imwrite(temp_mask_path, binary_mask)
    with open(temp_mask_path, "rb") as f:
        files = {"file": (temp_mask_path, f, "image/png")}
        data = {
            "image_height_mm": str(image_height_mm),
            "image_width_mm": str(image_width_mm)
        }
        resp = requests.post("http://127.0.0.1:8010/analyze_crack", files=files, data=data)
        if resp.status_code != 200:
            raise RuntimeError(f"裂缝分析失败: {resp.text}")
        result = resp.json()
    # 补全字段
    required_keys = ["A", "B", "C", "D", "倾角_deg", "倾向_deg", "走向_deg",
                     "裂缝宽度_FVA", "裂缝长度_FVTL", "裂缝密度_FVDC", "裂缝视孔隙度_FVPA"]
    for k in required_keys:
        if k not in result:
            result[k] = None
    return result

def call_vug_api(image_path, image_height_mm, image_width_mm, window_height_mm):
    with open(image_path, "rb") as f:
        files = {"file": (image_path, f, "image/png")}
        data = {
            "image_height_mm": str(image_height_mm),
            "image_width_mm": str(image_width_mm),
            "window_height_mm": str(window_height_mm)
        }
        resp = requests.post("http://127.0.0.1:8011/analyze_vug", files=files, data=data)
        if resp.status_code != 200:
            raise RuntimeError(f"孔洞分析失败: {resp.text}")
        return resp.json()
#五、滑窗分析（核心计算模块）
# ===== 滑窗分析 =====
# ===== 滑窗分析（总控调度器）=====
def sliding_window_analysis(
    image_path,
    model_ids,
    model_parameters,
    log_fn=None
):
    """
    总控滑窗分析函数
    - 根据 model_id 类型自动分派 fracture / vug
    - 统一参数入口 model_parameters
    - 统一结构化返回结果
    """

    results = {
        "fracture": [],
        "vug": []
    }

    # ---- 基础参数校验（工程必备）----
    required_base_keys = ["image_height_mm", "image_width_mm"]
    for k in required_base_keys:
        if k not in model_parameters:
            raise KeyError(f"model_parameters 缺少必要参数: {k}")
    #print(model_ids)
    print(type(model_ids))
    model_list = model_ids['model_ids']
    print(model_list)
    print(type(model_list))
    # ✅ 兼容 Executor 误传 dict 的情况
    for model_id in model_list:

        model_id_lower = model_id.lower()

        # ===== 裂缝模型 =====
        if "unet_fracture" in model_id_lower:

            if log_fn:
                log_fn(f"🧩 启动裂缝滑窗分析: {model_id}")

            out = sliding_window_unet_analysis(
                image_path=image_path,
                model_id=model_id,
                image_height_mm=model_parameters["image_height_mm"],
                image_width_mm=model_parameters["image_width_mm"],
                log_fn=log_fn
            )

            results["fracture"].append({
                "model_id": model_id,
                "mask": out.get("fracture_mask"),
                "metrics": out.get("fracture_metrics", [])
            })

        # ===== 孔洞模型 =====
        elif "unet_vug" in model_id_lower:

            if log_fn:
                log_fn(f"🧩 启动孔洞滑窗分析: {model_id}")

            mask_path, window_metrics, summary = sliding_window_vug_analysis(
                image_path=image_path,
                model_id=model_id,
                image_height_mm=model_parameters["image_height_mm"],
                image_width_mm=model_parameters["image_width_mm"],
                window_height_mm=model_parameters.get("window_height_mm", 1000),
                window_px=model_parameters.get("window_px"),
                log_fn=log_fn
            )

            results["vug"].append({
                "model_id": model_id,
                "mask": mask_path,
                "window_metrics": window_metrics,
                "summary": summary
            })

        else:
            if log_fn:
                log_fn(f"⚠️ 未识别的模型类型，已跳过: {model_id}")
    # ---------- 打印 mask_path 信息 ----------
    print("🧩 Fracture masks:")
    for f in results["fracture"]:
        print(f"- model_id: {f['model_id']}, mask_path: {f.get('mask')}")

    print("🟡 Vug masks:")
    for v in results["vug"]:
        print(f"- model_id: {v['model_id']}, mask_path: {v.get('mask')}")

    return results

def sliding_window_unet_analysis(image_path, model_id, image_height_mm, image_width_mm, log_fn=None):
    img = cv2.imread(image_path)
    H, W = img.shape[:2]
    window_px = W  # 可以自定义滑窗高度
    n_blocks = math.ceil(H / window_px)
    masks_full = np.zeros((H, W), dtype=np.uint8)
    curves_metrics_all = []

    for i in range(n_blocks):
        start_y = i * window_px
        end_y = min((i+1) * window_px, H)
        img_block = img[start_y:end_y, :]
        block_path = f"temp_block_{uuid.uuid4().hex}.png"
        #block_path = os.path.join(temp_dir, f"temp_block_{uuid.uuid4().hex}.png")
        cv2.imwrite(block_path, img_block)

        try:
            mask_result = call_unet_api(model_id, block_path)
            mask_block = cv2.imread(mask_result["mask"], cv2.IMREAD_GRAYSCALE)
            if mask_block is None:
                continue
            masks_full[start_y:end_y, :] = np.maximum(masks_full[start_y:end_y, :], mask_block)

            # 裂缝分析
            mask_path = preprocess_mask_for_analysis(mask_result["mask"], log_fn)
            single_masks = split_mask_to_contours(mask_path)
            for j, mask_single in enumerate(single_masks):
                temp_mask_path = f"temp_single_block_{i}_{j}.png"
                #temp_mask_path = os.path.join(temp_dir, f"temp_single_block_{i}_{j}.png")
                cv2.imwrite(temp_mask_path, mask_single)
                metrics = call_crack_api(temp_mask_path, image_height_mm, image_width_mm)
                metrics["y_offset"] = start_y  # ⚡ 添加块偏移
                curves_metrics_all.append(metrics)

        except Exception as e:
            if log_fn:
                log_fn(f"❌ 分块 U-Net 分析失败: {e}")

    mask_full_path = f"full_mask_{uuid.uuid4().hex}.png"
    #mask_full_path = os.path.join(temp_dir, f"full_mask_{uuid.uuid4().hex}.png")
    # ✅ 强制归一化为 0-255，保证显示效果与 SAM2 一致
    if masks_full.max() <= 1:
        masks_full = (masks_full * 255).astype(np.uint8)
    else:
        masks_full = masks_full.astype(np.uint8)

    cv2.imwrite(mask_full_path, masks_full)
    # ✅【关键修改】—— 返回 dict，严格匹配 plan.outputs
    fracture_metrics_filtered = [m for m in curves_metrics_all if
                                 all(k in m and m[k] is not None for k in ["A", "B", "C", "D"])]
    print(mask_full_path)
    return {
        "fracture_mask": mask_full_path,
        "fracture_metrics": fracture_metrics_filtered
    }

def sliding_window_vug_analysis(
    image_path,
    model_id,
    image_height_mm,
    image_width_mm,
    window_height_mm=1000,   # 默认 1 m
    window_px=None,          # ✅ 支持显式像素滑窗
    log_fn=None
):
    """
    孔洞统一分析流程（与裂缝滑窗保持一致）：
    - 内部滑窗：像素
    - 统计单位：mm
    - 返回：mask_full_path, window_metrics, summary
    """

    img = cv2.imread(image_path)
    if img is None:
        raise FileNotFoundError(f"未找到图像文件: {image_path}")

    H, W = img.shape[:2]

    # ===== px ↔ mm 映射 =====
    px_per_mm = H / image_height_mm
    mm_per_px = image_height_mm / H

    # ===== 滑窗优先级：window_px > window_height_mm =====
    if window_px is not None:
        window_px = int(window_px)
        window_height_mm = window_px * mm_per_px
    else:
        window_px = int(window_height_mm * px_per_mm)

    window_px = max(1, window_px)
    n_blocks = math.ceil(H / window_px)

    masks_full = np.zeros((H, W), dtype=np.uint8)

    window_vug_list = []
    total_vug_count = 0
    total_area_mm2 = 0
    all_CVPA, all_CDENS, all_CSIZE = [], [], []

    for i in range(n_blocks):
        start_y = i * window_px
        end_y = min(start_y + window_px, H)

        depth_start_mm = start_y * mm_per_px
        depth_end_mm = end_y * mm_per_px
        window_depth_mm = depth_end_mm - depth_start_mm

        img_block = img[start_y:end_y, :]
        block_path = f"temp_vug_block_{uuid.uuid4().hex}.png"
        cv2.imwrite(block_path, img_block)

        try:
            # ---------- ① U-Net 孔洞分割 ----------
            mask_result = call_unet_api(model_id, block_path)
            mask_block = cv2.imread(mask_result["mask"], cv2.IMREAD_GRAYSCALE)
            if mask_block is None:
                continue

            masks_full[start_y:end_y, :] = np.maximum(
                masks_full[start_y:end_y, :],
                mask_block
            )

            # ---------- ② 单滑窗孔洞统计 ----------
            cleaned_mask = preprocess_mask_for_analysis(mask_result["mask"], log_fn)

            result = call_vug_api(
                cleaned_mask,
                image_height_mm=window_depth_mm,
                image_width_mm=image_width_mm,
                window_height_mm=window_depth_mm
            )

            summary = result.get("summary", {})
            vugs = result.get("vugs", [])

            window_info = {
                "depth_start_mm": round(depth_start_mm, 2),
                "depth_end_mm": round(depth_end_mm, 2),
                "vug_count": summary.get("vug_count", 0),
                "total_area_mm2": summary.get("total_area_mm2", 0),
                "CVPA": summary.get("CVPA", 0),
                "CDENS": summary.get("CDENS", 0),
                "CSIZE": summary.get("CSIZE", 0),
                "vugs": vugs
            }

            window_vug_list.append(window_info)

            total_vug_count += window_info["vug_count"]
            total_area_mm2 += window_info["total_area_mm2"]
            all_CVPA.append(window_info["CVPA"])
            all_CDENS.append(window_info["CDENS"])
            all_CSIZE.append(window_info["CSIZE"])

            if log_fn:
                log_fn(f"✅ 孔洞滑窗 {i+1}/{n_blocks} 完成")

        except Exception as e:
            if log_fn:
                log_fn(f"⚠️ 分块孔洞分析失败: {e}")

    # ---------- 保存整图 mask ----------
    mask_full_path = f"full_mask_vug_{uuid.uuid4().hex}.png"
    if masks_full.max() <= 1:
        masks_full = (masks_full * 255).astype(np.uint8)
    else:
        masks_full = masks_full.astype(np.uint8)
    cv2.imwrite(mask_full_path, masks_full)

    # ---------- 整井汇总 ----------
    summary_metrics = {
        "total_vug_count": total_vug_count,
        "total_area_mm2": total_area_mm2,
        "mean_CVPA": np.mean(all_CVPA) if all_CVPA else 0,
        "mean_CDENS": np.mean(all_CDENS) if all_CDENS else 0,
        "mean_CSIZE": np.mean(all_CSIZE) if all_CSIZE else 0
    }

    if log_fn:
        log_fn(
            f"📊 孔洞滑窗分析完成: "
            f"总孔洞数={total_vug_count}, 总面积={total_area_mm2:.2f} mm²"
        )
    print(mask_full_path)
    return mask_full_path, window_vug_list, summary_metrics




# ===== SAM2 滑窗分析 =====
def sliding_window_sam2_analysis(image_path, image_height_mm, image_width_mm, log_fn=None):
    """
    滑窗调用 SAM2（prompt_free 模式），并拼接整体掩码
    """
    img = cv2.imread(image_path)
    H, W = img.shape[:2]
    window_px = W  # 与 U-Net 一致，纵向滑窗
    n_blocks = math.ceil(H / window_px)
    masks_full = np.zeros((H, W), dtype=np.uint8)
    curves_metrics_all = []

    for i in range(n_blocks):
        start_y = i * window_px
        end_y = min((i + 1) * window_px, H)
        img_block = img[start_y:end_y, :]
        block_path = f"temp_sam2_block_{uuid.uuid4().hex}.png"
        cv2.imwrite(block_path, img_block)

        try:
            # prompt_free 模式分割
            sam2_result = call_sam2_api(block_path, prompt_type="prompt_free")
            mask_b64 = sam2_result.get("mask", None)
            if not mask_b64:
                if log_fn:
                    log_fn(f"⚠️ SAM2 第 {i+1}/{n_blocks} 块未返回有效 mask")
                continue

            mask_bytes = base64.b64decode(mask_b64)
            mask_block_path = f"sam2_block_mask_{uuid.uuid4().hex}.png"
            #mask_block_path = os.path.join(temp_dir, f"sam2_block_mask_{uuid.uuid4().hex}.png")
            with open(mask_block_path, "wb") as f:
                f.write(mask_bytes)

            mask_block = cv2.imread(mask_block_path, cv2.IMREAD_GRAYSCALE)
            if mask_block is None:
                continue

            masks_full[start_y:end_y, :] = np.maximum(
                masks_full[start_y:end_y, :], mask_block
            )

            # 裂缝参数分析（同 U-Net）
            mask_path = preprocess_mask_for_analysis(mask_block_path, log_fn)
            single_masks = split_mask_to_contours(mask_path)
            for j, mask_single in enumerate(single_masks):
                temp_mask_path = f"temp_single_sam2_block_{i}_{j}.png"
                #temp_mask_path = os.path.join(temp_dir, f"temp_single_sam2_block_{i}_{j}.png")
                cv2.imwrite(temp_mask_path, mask_single)
                metrics = call_crack_api(temp_mask_path, image_height_mm, image_width_mm)
                metrics["y_offset"] = start_y
                curves_metrics_all.append(metrics)

            if log_fn:
                log_fn(f"✅ SAM2 滑窗块 {i+1}/{n_blocks} 分析完成，裂缝 {len(single_masks)} 条")

        except Exception as e:
            if log_fn:
                log_fn(f"❌ SAM2 滑窗块 {i+1}/{n_blocks} 失败: {e}")

    mask_full_path = f"full_sam2_mask_{uuid.uuid4().hex}.png"
    #mask_full_path = os.path.join(temp_dir, f"full_sam2_mask_{uuid.uuid4().hex}.png")
    cv2.imwrite(mask_full_path, masks_full)
    return mask_full_path, curves_metrics_all
#六、结果可视化与绘制
# ===== 绘制最终结果（保持不变，使用 metrics["y_offset"]） =====
# ========== 绘制蝌蚪图 (改进版) ==========

def extract_masks_from_sliding_results(sliding_results, target_classes=None, log_fn=None):
    """
    从 sliding_results 中提取 mask 路径
    - sliding_results: {"fracture": [...], "vug": [...]}
    - target_classes: ["fracture", "vug"] 默认提取全部
    返回: list of dicts {"mask": mask_path}, list of model_ids
    """
    if target_classes is None:
        target_classes = sliding_results.keys()

    masks_list = []
    model_ids = []

    for cls in target_classes:
        items = sliding_results.get(cls, [])
        for item in items:
            mask_path = item.get("mask") or item.get("mask_path")  # 兼容 fracture/vug
            if not mask_path:
                if log_fn:
                    log_fn(f"⚠️ {cls} 条目缺少 mask 路径: {item}")
                continue
            masks_list.append({"mask": mask_path})
            model_ids.append(f"unet_{cls}")

    return masks_list, model_ids



def plot_tadpole_from_crack_results_v2(valid_curves,
                                   save_path="tadpole.png",
                                   image_height_px=1422,
                                   image_width_px=472,
                                   dpi=100,
                                   start_depth=1170000,
                                   end_depth=1172500):
    """
    蝌蚪图绘制（直接像素坐标绘制 + 深度刻度显示）

    参数：
    - valid_curves: list[dict]，每项包含 "D", "y_offset", "倾角_deg", "倾向_deg"
    - save_path: 输出路径
    - image_height_px, image_width_px: 图像像素尺寸
    - dpi: 分辨率
    - start_depth, end_depth: y轴显示深度范围
    """
    if not valid_curves:
        raise ValueError("没有有效曲线数据，无法生成蝌蚪图")

    # ==== 图像尺寸 ====
    fig_w_in = image_width_px / dpi
    fig_h_in = image_height_px / dpi
    fig, ax = plt.subplots(figsize=(fig_w_in, fig_h_in), dpi=dpi)

    # ==== 坐标轴设置 ====
    ax.set_xlim(0, 90)               # 倾角范围
    ax.set_ylim(image_height_px, 0)  # 上浅下深
    ax.set_xlabel("Dip Angle (°)", fontsize=9)
    ax.set_ylabel("Depth (mm)", fontsize=9)

    # y轴刻度显示深度
    def depth_formatter(y_px, pos):
        return f"{start_depth + (y_px / image_height_px) * (end_depth - start_depth):.0f}"
    ax.yaxis.set_major_formatter(FuncFormatter(depth_formatter))
    ax.xaxis.set_ticks_position("top")
    ax.xaxis.set_label_position("top")
    ax.grid(alpha=0.3)

    # ==== 绘制蝌蚪 ====
    cmap = plt.get_cmap("tab10")
    tail_len_px = 40  # 尾巴长度（像素）

    print("纵向位置(y_offset, D):", [(m.get("y_offset",0), m.get("D",0)) for m in valid_curves])

    for i, m in enumerate(valid_curves):
        dip = m.get("倾角_deg", 0)
        az = m.get("倾向_deg", 0)
        y_px = m.get("y_offset", 0) + m.get("D", 0)
        color = cmap(i % 10)

        # 头部
        ax.scatter(dip, y_px, color=color, s=30, zorder=3)

        # 尾巴方向
        dx = tail_len_px * np.sin(np.deg2rad(az))
        dy = tail_len_px * np.cos(np.deg2rad(az))
        ax.arrow(dip, y_px, dx/90*10, dy, color=color, alpha=0.8,
                 width=0.5, head_width=2.5, length_includes_head=True)

    plt.subplots_adjust(left=0.2, right=0.95, top=0.92, bottom=0.08)
    fig.savefig(save_path, dpi=dpi)
    plt.close(fig)
    print(f"✅ 蝌蚪图保存: {save_path} ({image_width_px}×{image_height_px}px)")

def create_tadpole_overlay(valid_curves,
                           overlay_img_path="temp_input_overlay_final.png",
                           save_path="tadpole_overlay.png",
                           dpi=100,
                           depth_start=1170000,
                           depth_end=1172500,
                           tail_len_px=40):
    """
    绘制蝌蚪图并与原始图横向拼接，左侧添加深度道。

    参数：
    - valid_curves: list[dict] 包含 "D", "y_offset", "倾角_deg", "倾向_deg"
    - overlay_img_path: 原始图路径
    - save_path: 输出路径
    - dpi: 分辨率
    - depth_start, depth_end: y轴深度范围
    - tail_len_px: 蝌蚪尾巴长度
    """
    # === 读取原图 ===
    base_img = cv2.imread(overlay_img_path)
    img_h, img_w = base_img.shape[:2]

    # === 蝌蚪图绘制 ===
    fig, ax = plt.subplots(figsize=(img_w/dpi, img_h/dpi), dpi=dpi)

    ax.set_xlim(0, 90)
    ax.set_ylim(img_h, 0)  # 上浅下深
    ax.axis('off')  # 去掉多余坐标

    cmap = plt.get_cmap("tab10")
    for i, m in enumerate(valid_curves):
        dip = m.get("倾角_deg", 0)
        az = m.get("倾向_deg", 0)
        y_px = m.get("y_offset", 0) + m.get("D", 0)
        color = cmap(i % 10)
        ax.scatter(dip, y_px, color=color, s=30, zorder=3)
        dx = tail_len_px * np.sin(np.deg2rad(az))
        dy = tail_len_px * np.cos(np.deg2rad(az))
        ax.arrow(dip, y_px, dx/90*10, dy, color=color, alpha=0.8,
                 width=0.5, head_width=2.5, length_includes_head=True)

    fig.tight_layout(pad=0)
    fig.canvas.draw()

    # 将绘图保存到 numpy 数组
    tadpole_img = np.frombuffer(fig.canvas.tostring_rgb(), dtype=np.uint8)
    tadpole_img = tadpole_img.reshape(fig.canvas.get_width_height()[::-1] + (3,))
    plt.close(fig)

    # === 添加深度道 ===
    depth_width = 60
    depth_img = np.zeros((img_h, depth_width, 3), dtype=np.uint8) + 255  # 白色背景
    n_ticks = 10
    for i in range(n_ticks+1):
        y = int(i * img_h / n_ticks)
        depth_val = int(depth_start + (y/img_h)*(depth_end-depth_start))
        cv2.putText(depth_img, f"{depth_val}", (2, y+5),
                    cv2.FONT_HERSHEY_SIMPLEX, 0.4, (0,0,0), 1, cv2.LINE_AA)

    # === 横向拼接：深度道 + 原图 + 蝌蚪图 ===
    tadpole_img_resized = cv2.resize(tadpole_img, (img_w, img_h))
    final_img = np.hstack([depth_img, base_img, tadpole_img_resized])

    # === 保存 ===
    cv2.imwrite(save_path, final_img)
    print(f"✅ 拼接蝌蚪图保存: {save_path}")
def draw_final_results(base_image_path, unet_results, yolo_results, H):
    base_img = Image.open(base_image_path).convert("RGB")
    img_np = np.array(base_img)
    draw = ImageDraw.Draw(base_img)

    # ---------- YOLO 框 ----------
    if isinstance(yolo_results, dict):
        detections = yolo_results.get("detections", [])
    else:
        detections = []  # 如果不是 dict，就忽略
    for det in detections:
        cls = det.get("class", "")
        conf = det.get("confidence", 0)
        bbox = list(map(int, det.get("bbox", [])))
        draw.rectangle(bbox, outline="red", width=2)
        draw.text((bbox[0], max(0, bbox[1]-12)), f"{cls} {conf:.2f}", fill="red")

    # ---------- U-Net / SAM2 掩码 + 裂缝曲线 ----------
    if isinstance(unet_results, dict):
        items = []
        for v in unet_results.values():
            if isinstance(v, list):
                items.extend(v)
        unet_results = items
    elif not isinstance(unet_results, list):
        unet_results = []

    for item in unet_results:
        if not isinstance(item, dict):
            continue  # 跳过非 dict
        cls = str(item.get("class", "")).lower()
        print("🧬 drawing item class =", cls)
        mask_entry = item.get("mask_result")
        if isinstance(mask_entry, dict):
            mask_path = mask_entry.get("mask")
        elif isinstance(mask_entry, str):
            mask_path = mask_entry
        else:
            mask_path = None

        if not mask_path or not os.path.exists(mask_path):
            continue

        mask_img = cv2.imread(mask_path, cv2.IMREAD_GRAYSCALE)
        if mask_img is None:
            continue

        # ---------- 掩码轮廓 ----------
        contours, _ = cv2.findContours((mask_img>0).astype(np.uint8), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        for contour in contours:
            pts = [tuple(pt[0]) for pt in contour]
            draw.line(pts, fill=(0, 255, 0), width=2)

        # ---------- 裂缝曲线 ----------
        if cls == "fracture":
            metrics_list = item.get("metrics_list", [])
            print(metrics_list)
            print("📐 metrics_list length =", len(metrics_list))
            if not isinstance(metrics_list, list):
                continue
            for metrics in metrics_list:
                if all(k in metrics for k in ["A", "B", "C", "D"]):
                    A = metrics["A"]
                    B = metrics["B"]
                    C = metrics["C"]
                    D = metrics["D"]
                    y_offset = metrics.get("y_offset", 0)
                    if None not in [A, B, C, D]:
                        w = img_np.shape[1]
                        x_fit = np.arange(w)
                        y_fit = A * np.sin(B * x_fit + C) + D + y_offset
                        points = [(int(x), int(y)) for x, y in zip(x_fit, y_fit) if 0 <= int(y) < H]
                        print(
                            f"📈 curve y range = [{np.min(y_fit):.1f}, {np.max(y_fit):.1f}], H = {H}"
                        )
                        draw.line([(0, H // 2), (img_np.shape[1], H // 2)], fill=(255, 0, 0), width=3)
                        for i in range(len(points)-1):
                            draw.line([points[i], points[i+1]], fill=(0, 0, 0), width=2)

    out_path = base_image_path.replace(".png","_final.png")
    base_img.save(out_path)
    print(out_path)
    return out_path

# 记录
def record_execution_state(
    intent: str = "",
    planner: Dict[str, Any] = None,
    flags: Dict[str, Any] = None,
    raw_sliding: Any = None,
    refined_curves: Any = None,
    analysis_log: Any = None,
    extra: Dict[str, Any] = None,
    output_dir: str = "records",
    prefix: str = "execution_record"
):
    """
    📌 Recorder Tool

    Parameters
    ----------
    intent : str
        Planner intent or task description
    planner : dict
        Planner-related outputs (selected_models, parameters, etc.)
    flags : dict
        Runtime flags (enable_sam2, enable_reflection, etc.)
    raw_sliding : any
        Raw sliding_window_analysis results
    refined_curves : any
        Post-reflection curves (if exists)
    analysis_log : any
        DeepSeek reflection logs
    extra : dict
        Any additional user-defined content
    output_dir : str
        Directory to store records
    prefix : str
        Filename prefix

    Returns
    -------
    dict
        {
            "record_path": "...",
            "record_id": "...",
            "timestamp": "..."
        }
    """

    os.makedirs(output_dir, exist_ok=True)

    record_id = uuid.uuid4().hex
    timestamp = datetime.datetime.now().isoformat()

    record = {
        "record_id": record_id,
        "timestamp": timestamp,
        "intent": intent,
        "planner": planner or {},
        "flags": flags or {},
        "raw_sliding_results": raw_sliding,
        "refined_curves": refined_curves,
        "analysis_log": analysis_log,
        "extra": extra or {}
    }

    record_path = os.path.join(
        output_dir,
        f"{prefix}_{record_id}.json"
    )

    with open(record_path, "w", encoding="utf-8") as f:
        json.dump(record, f, indent=2, ensure_ascii=False)

    return {
        "record_path": record_path,
        "record_id": record_id,
        "timestamp": timestamp
    }
# 派生变量
def generate_x_points(curves_metrics, image_width_px):
    """
    Generate x_points for each curve.
    """
    if curves_metrics is None:
        return []
    valid_curves = [m for m in curves_metrics if all(
        k in m and m[k] is not None for k in ["A", "B", "C", "D"]
    )]
    x_points_list = []
    x_points_list=[np.arange(image_width_px) for _ in valid_curves]
    return x_points_list
def extract_fracture_metrics(sliding_results: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Extract and flatten fracture metrics from sliding_results.

    Returns:
        metrics_list: List[dict]
            Each dict corresponds to ONE fracture curve, with metadata.
    """

    if not isinstance(sliding_results, dict):
        raise ValueError("sliding_results must be a dict")

    fracture_blocks = sliding_results.get("fracture", [])
    if not fracture_blocks:
        return []

    metrics_list: List[Dict[str, Any]] = []

    for block_idx, block in enumerate(fracture_blocks):

        if not isinstance(block, dict):
            continue

        model_id = block.get("model_id", "unknown")
        metrics = block.get("metrics", [])

        # metrics 本身是 list
        if not isinstance(metrics, list):
            continue

        for curve_idx, curve_metrics in enumerate(metrics):
            if not isinstance(curve_metrics, dict):
                continue

            # ✅ 复制一份，避免原地污染
            item = dict(curve_metrics)

            # ✅ 注入元信息（非常重要）
            item["_model_id"] = model_id
            item["_fracture_block_index"] = block_idx
            item["_curve_index"] = curve_idx

            metrics_list.append(item)

    return metrics_list
import numpy as np
from typing import Dict, Any, List, Tuple

def extract_metrics_and_xpoints(
    sliding_results: Dict[str, Any],
    image_width_px: int
) -> Dict[str, Any]:
    """
    Extract and flatten fracture metrics from sliding_results
    and generate x_points for each curve.

    Returns:
        dict with keys:
            - curves_metrics
            - x_points_list
    """

    if not isinstance(sliding_results, dict):
        raise ValueError("sliding_results must be a dict")
    if not isinstance(image_width_px, int) or image_width_px <= 0:
        raise ValueError("image_width_px must be a positive integer")

    fracture_blocks = sliding_results.get("fracture", [])
    if not fracture_blocks:
        return {"curves_metrics": [], "x_points_list": []}

    metrics_list: List[Dict[str, Any]] = []
    x_points_list: List[np.ndarray] = []

    for block_idx, block in enumerate(fracture_blocks):
        if not isinstance(block, dict):
            continue

        model_id = block.get("model_id", "unknown")
        metrics = block.get("metrics", [])
        if not isinstance(metrics, list):
            continue

        for curve_idx, curve_metrics in enumerate(metrics):
            if not isinstance(curve_metrics, dict):
                continue

            # 复制一份，避免污染
            item = dict(curve_metrics)
            # 注入元信息
            item["_model_id"] = model_id
            item["_fracture_block_index"] = block_idx
            item["_curve_index"] = curve_idx
            metrics_list.append(item)

            # 生成 x_points
            if all(k in item and item[k] is not None for k in ["A", "B", "C", "D"]):
                x_points_list.append(np.arange(image_width_px))

    return {"curves_metrics": metrics_list, "x_points_list": x_points_list}
# ====== overwrite_metrics 工具函数 ======
def overwrite_metrics(target, source, flags=None, log_fn=None):
    """
    条件覆盖 metrics 列表
    - target: 原始 metrics 列表（会被覆盖）
    - source: 新的 metrics 列表
    - flags: 可选字典，用于判断是否启用覆盖逻辑
    - log_fn: 可选日志函数
    """
    enable_reflection = False
    if flags and isinstance(flags, dict):
        enable_reflection = flags.get("enable_reflection", False)

    if enable_reflection:
        if log_fn:
            log_fn(f"🔄 overwrite_metrics: 覆盖 {len(target)} 条 metrics 为 {len(source)} 条 metrics")
        # ✅ 返回覆盖后的列表，保持 Executor 需要的 key
        return {
            "derived.fracture.metrics_list": source
        }
    else:
        if log_fn:
            log_fn("⚠️ overwrite_metrics: 条件未满足，保持原 metrics")
        return {
            "derived.fracture.metrics_list": target
        }
# 解析工具
def resolve_masks_from_sliding_results(sliding_results, log_fn=None):
    """
    将 sliding_results 中的 fracture / vug mask 显式解析出来
    用于 overlay_masks 之前
    """
    masks = []

    # -------- fracture --------
    fracture = sliding_results.get("fracture")
    if isinstance(fracture, dict):
        mask = fracture.get("mask")
        if mask:
            masks.append(mask)

    # -------- vug --------
    vug = sliding_results.get("vug")
    if isinstance(vug, dict):
        mask = vug.get("mask")
        if mask:
            masks.append(mask)

    if log_fn:
        log_fn(f"✅ resolved masks: {masks}")

    return masks
def extract_overlay_masks(sliding_results, log_fn=None):
    """
    Executor 规则：
    - outputs 是完整路径 → return dict 的 key 必须完全一致
    """

    fracture_mask = None
    vug_mask = None

    if "fracture" in sliding_results and sliding_results["fracture"]:
        fracture_mask = sliding_results["fracture"][0].get("mask")

    if "vug" in sliding_results and sliding_results["vug"]:
        vug_mask = sliding_results["vug"][0].get("mask")

    if log_fn:
        log_fn(
            f"✅ extract_overlay_masks: "
            f"fracture={fracture_mask}, vug={vug_mask}"
        )
    print(vug_mask)
    print(fracture_mask)
    # ✅ 关键修复点（完全匹配 outputs）
    return {
        "derived.overlay.fracture_mask": fracture_mask,
        "derived.overlay.vug_mask": vug_mask
    }
def rebuild_unet_results(sliding_results, fracture_metrics_list, log_fn=None):
    """
    将 DAG 产出的 sliding_results + fracture_metrics_list
    重建为 draw_final_results 期望的 unet_results 结构
    """

    # 🔧 关键修复：如果 metrics 被包了一层 dict，直接拆
    if isinstance(fracture_metrics_list, dict):
        if len(fracture_metrics_list) == 1:
            fracture_metrics_list = list(fracture_metrics_list.values())[0]

    unet_results = []

    for item in sliding_results.get("fracture", []):
        unet_results.append({
            "class": "fracture",
            "mask_result": {"mask": item.get("mask")},
            "metrics_list": fracture_metrics_list   # ✅ 现在是 list[dict]
        })

    if log_fn:
        log_fn(f"🧱 unet_results: {len(unet_results)} items rebuilt")

    return {
        "unet_results": unet_results
    }
def build_final_result_json(
    user_prompt,
    yolo_results,
    sam2_results,
    unet_results,
    sliding_results,
    params_used,
    log_fn=None
):
    import datetime
    # 假设 sliding_window_analysis 返回结果已经保存在 vug_results 里
    #print("原始 vug_results:", vug_results)

    # ✅ 正确拿出孔洞列表

    vug_results_list = sliding_results.get("vug", [])
    #print("处理后的 vug_results_list:", vug_results_list)

    result = {
        "user_prompt": user_prompt,
        "yolo_result": yolo_results,
        "sam2_results": sam2_results,
        "unet_results": unet_results,
        "vug_results": vug_results_list,
        "params_used": params_used,
        "timestamp": datetime.datetime.now().isoformat()
    }

    if log_fn:
        log_fn("📦 Final result JSON assembled")
    #print(result)
    return {
        "deepseek_json": result
    }
