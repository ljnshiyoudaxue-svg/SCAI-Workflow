import os
import shutil
import json
import numpy as np
import datetime
import cv2
from yolo_agent2 import call_yolo_api
from PIL import Image
import Agent_tools
from Agent_tools import parse_deepseek_json
from Agent_tools import sliding_window_vug_analysis
import requests
# ✅ 导入核心 API
from deepseek_agent_eval import (
    sliding_window_sam2_analysis,
    sliding_window_unet_analysis,
    overlay_masks,
    draw_final_results
)
from deepseek_agent_eval2_localllm import (
    deepseek_decide_models,
    generate_prompt,
    deepseek_filter_curves_safe
)






import time

def segment_image_TCA(
        user_prompt,
        image_file,
        image_height_mm=2500,
        image_width_mm=215,
        enable_sam2=True,
        enable_reflection=True,
        enable_prompt_mode="default",
        multi_model=True,
        clean_temp=True
):
    """
    🎯 Task-Controlled Agent (TCA)
    自动执行裂缝/孔洞分割 + 模型决策 + SAM2滑窗 + 自反思复核 + 报告结果整合
    支持四种提示词模式: default, zero-shot, few-shot, hard-constraint
    """
    logs = []

    def log(msg):
        logs.append(msg)
        print(msg)

    # =============== 🧩 0️⃣ 初始化临时目录 ===============
    base_name = os.path.splitext(os.path.basename(getattr(image_file, "filename", "input.png")))[0]
    temp_dir = os.path.join("temp_workspace_TCA", base_name)
    os.makedirs(temp_dir, exist_ok=True)
    log(f"📁 [Context] 临时工作目录: {temp_dir}")

    # =============== 🖼️ 1️⃣ 保存输入图像 ===============
    image_path = os.path.join(temp_dir, "input.png")
    try:
        if hasattr(image_file, "save"):
            image_file.save(image_path)
        else:
            shutil.copy(image_file, image_path)
        log(f"📄 [Context] 输入图像: {image_path}")
    except Exception as e:
        log(f"❌ 输入图像保存失败: {e}")
        return None, [], None, None, "\n".join(logs)

    # =============== 🧠 2️⃣ YOLO 检测 ===============
    log("🚀 执行 Action: YOLO 检测")
    try:
        yolo_results = call_yolo_api(image_path)
        log(f"✅ YOLO 检测完成，共 {len(yolo_results.get('detections', []))} 个目标")
    except Exception as e:
        yolo_results = {"detections": []}
        log(f"⚠️ YOLO 检测失败: {e}")

    # =============== 💬 3️⃣ DeepSeek 模型决策 ===============
    log("🚀 执行 Action: 模型决策")
    model_ids = []
    user_prompt = "请分析这张电成像裂缝"
    user_prompt_safe = user_prompt
    user_prompt_final = f"[PROMPT_MODE={enable_prompt_mode}] {user_prompt_safe}" if enable_prompt_mode != "default" else user_prompt_safe

    try:
        t0 = time.time()  # ⏱️ 开始计时
        model_ids, parameters = deepseek_decide_models(
            user_prompt_final,
            yolo_results,
            enable_prompt_mode  # ✅ 四种模式传入
        )
        t_decide = time.time() - t0  # ⏱️ 耗时
        log(f"🕒 模型决策耗时: {t_decide:.2f} 秒")

        if not multi_model:
            model_ids = model_ids[:1]
        log(f"✅ 模型决策完成: {model_ids}")
    except Exception as e:
        model_ids = []
        log(f"⚠️ 模型决策失败: {e}")

    # =============== 🧩 4️⃣ SAM2 分割 ===============
    sam2_results = []
    sam2_mask_path = None
    if enable_sam2:
        log("🚀 执行 Action: SAM2 分割")
        try:
            sam2_mask_path, sam2_metrics = sliding_window_sam2_analysis(
                image_path, image_height_mm, image_width_mm, log_fn=log
            )
            if sam2_mask_path:
                new_sam2_path = os.path.join(temp_dir, "mask_sam2.png")
                shutil.move(sam2_mask_path, new_sam2_path)
                sam2_mask_path = new_sam2_path
            valid_curves = [m for m in sam2_metrics if all(k in m and m[k] is not None for k in ["A", "B", "C", "D"])]
            sam2_results = [{"mask_result": {"mask": sam2_mask_path}, "metrics_list": valid_curves}]
            log(f"✅ SAM2 分割完成，有效曲线: {len(valid_curves)}")
        except Exception as e:
            log(f"❌ SAM2 分割失败: {e}")
    else:
        log("⚙️ 已禁用 SAM2 分割")

    # =============== 🧬 5️⃣ U-Net 分割 + 自反思复核 ===============
    unet_results = []
    vug_results = None

    try:
        img = cv2.imread(image_path)
        image_height_px, image_width_px = img.shape[:2]
    except Exception as e:
        image_height_px, image_width_px = 0, 0
        log(f"⚠️ 无法读取图像尺寸: {e}")

    # ---------- 裂缝/孔洞分析循环 ----------
    print(model_ids)
    print(type(model_ids))
    for model_id in model_ids:
        # ---------- 裂缝分析 ----------
        if model_id.lower() == "unet_fracture":
            try:
                log(f"📌 [Action] 调用裂缝 U-Net 模型: {model_id}")
                mask_path, metrics_list = sliding_window_unet_analysis(
                    image_path, model_id, image_height_mm, image_width_mm, log_fn=log
                )

                if mask_path:
                    new_mask_path = os.path.join(temp_dir, f"mask_unet_{model_id}.png")
                    shutil.move(mask_path, new_mask_path)
                    mask_path = new_mask_path

                valid_curves = [m for m in metrics_list if
                                all(k in m and m[k] is not None for k in ["A", "B", "C", "D"])]

                # ✅ 自反思复核
                if enable_reflection and valid_curves:
                    log("📌 启用 DeepSeek 自反思复核...")
                    t1 = time.time()
                    x_points = [np.arange(image_width_px) for _ in valid_curves]
                    curves_filtered, analysis_log, _ = deepseek_filter_curves_safe(
                        valid_curves,
                        x_points,
                        image_height_px=image_height_px,
                        image_width_px=image_width_px,
                        log_fn=log,
                        strategy=enable_prompt_mode
                    )
                    t_filter = time.time() - t1
                    log(f"🕒 自反思复核耗时: {t_filter:.2f} 秒")
                    valid_curves = curves_filtered
                    for entry in analysis_log:
                        log(f"    {entry}")

                unet_results.append({
                    "class": model_id.split("_")[1] if "_" in model_id else model_id,
                    "mask_result": {"mask": mask_path},
                    "metrics_list": valid_curves
                })
                log(f"✅ 裂缝模型 {model_id} 完成, 有效曲线: {len(valid_curves)}")

            except Exception as e:
                log(f"❌ 裂缝模型 {model_id} 分析失败: {e}")

        # ---------- 孔洞分析 ----------
        elif model_id.lower() == "unet_vug":
            try:
                log(f"📌 调用孔洞 U-Net 模型: {model_id}")
                mask_path, vug_window_metrics, vug_summary = sliding_window_vug_analysis(
                    image_path=image_path,
                    model_id=model_id,
                    image_height_mm=image_height_mm,
                    image_width_mm=image_width_mm,
                    window_px=472,
                    log_fn=log
                )
                log(f"✅ 孔洞分析完成: 总孔洞数={vug_summary['total_vug_count']}, 总面积={vug_summary['total_area_mm2']:.2f} mm²")
                vug_results = {
                    "window_metrics": vug_window_metrics,
                    "summary": vug_summary,
                    "mask_path": mask_path
                }

            except Exception as e:
                log(f"⚠️ 孔洞模型 {model_id} 分析失败: {e}")

    # =============== 🧾 6️⃣ 绘制结果与报告 ===============
    try:
        all_masks = []
        all_model_ids = []

        # ---------- U-Net 裂缝 ----------
        for r in unet_results:
            all_masks.append(r["mask_result"])
            all_model_ids.append(f"unet_{r['class']}")

        # ---------- U-Net 孔洞 ----------
        if vug_results and "mask_path" in vug_results:
            all_masks.append({"mask": vug_results["mask_path"]})
            all_model_ids.append("unet_vug")

        # ---------- SAM2 ----------
        for r in sam2_results:
            all_masks.append(r["mask_result"])
            all_model_ids.append("sam2_prompt_free")
        overlay_path = overlay_masks(all_masks, all_model_ids, image_path)

        final_overlay_path = draw_final_results(
            overlay_path, unet_results, yolo_results, H=image_height_px
        )
        log(f"✅ 最终叠加图生成: {final_overlay_path}")
    except Exception as e:
        final_overlay_path = None
        log(f"❌ 报告生成失败: {e}")
    # =============== 🧩 7️⃣ 保存 JSON ===============
    result = {
        "user_prompt": user_prompt_final,
        "yolo_result": yolo_results,
        "sam2_results": sam2_results,
        "unet_results": unet_results,
        "vug_results": vug_results,
        "params_used": {
            "enable_sam2": enable_sam2,
            "enable_prompt_mode": enable_prompt_mode,
            "multi_model": multi_model
        },
        "timestamp": datetime.datetime.now().isoformat()
    }
    # ===============================
    # 🧾 统一报告生成（与 segment_image_gradio 一致）
    # ===============================

    try:
        deepseek_result_json = parse_deepseek_json(result, 2, 5387)
        report_api_url = "http://127.0.0.1:9095/generate_comprehensive_report"
        payload = {
            "result": deepseek_result_json,  # 第二轮修正后的 JSON
            "image_path": image_path,  # 复核图像路径
            "image_url": image_path
        }

        log("🔹 调用报告生成接口 ...")
        report_resp = requests.post(report_api_url, json=payload, timeout=120)

        if report_resp.status_code == 200:
            try:
                resp_json = report_resp.json()
                report_path = resp_json.get("report_path", "")
                report_preview = resp_json.get("report_preview", "")
                log("✅ 报告生成完成")
                log(f"📄 报告路径: {report_path}")
                log(f"📋 报告预览: {report_preview[:500]}...")
            except Exception:
                # 若不是 JSON 格式（例如返回 Markdown），则写入临时 Word
                log("⚠️ 报告接口返回非 JSON 格式，尝试写入 Word ...")
                report_text = report_resp.text.strip()

                from docx import Document
                from docx.shared import Inches
                doc = Document()
                doc.add_heading("电成像综合地质分析报告", level=1)
                for line in report_text.split("\n"):
                    if line.startswith("#"):
                        doc.add_heading(line.replace("#", "").strip(), level=line.count("#"))
                    elif line.startswith("- "):
                        doc.add_paragraph(line[2:], style="List Bullet")
                    else:
                        doc.add_paragraph(line.strip())
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(5.5))
                report_path = f"DeepSeek_Report_TCA3.docx"
                doc.save(report_path)
                log(f"✅ 报告已写入 Word 文件: {report_path}")
        else:
            log(f"❌ 报告生成失败: HTTP {report_resp.status_code}")

    except Exception as e:
        log(f"⚠️ 报告生成时出现异常: {e}")

    # 最后更新状态
    result_json_path = os.path.join(temp_dir, "tca_result.json")
    with open(result_json_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    log(f"📊 结果 JSON 已保存: {result_json_path}")

    # =============== 🧹 8️⃣ 清理临时目录 ===============
    if clean_temp:
        shutil.rmtree(temp_dir, ignore_errors=True)
        log(f"🧹 临时目录已清理: {temp_dir}")

    log("✅ 当前任务完成")
    return sam2_mask_path, unet_results, final_overlay_path, result_json_path, "\n".join(logs)



