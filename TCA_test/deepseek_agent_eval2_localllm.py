import tempfile
import threading
import base64
import requests
from PIL import Image, ImageDraw
from yolo_agent2 import call_yolo_api  # YOLO接口
import uuid
import matplotlib
matplotlib.use("Agg")
import math
from collections import Counter
import matplotlib.pyplot as plt
from matplotlib.ticker import FuncFormatter
from openai import OpenAI
import datetime
import cv2
import numpy as np
import os
import cv2
import uuid
import json
import shutil
import numpy as np
import datetime
import Agent_tools
from Agent_tools import call_vug_api,overlay_masks,sliding_window_unet_analysis,sliding_window_vug_analysis,sliding_window_sam2_analysis,draw_final_results
from Agent_tools import call_unet_api,save_base64_mask,call_sam2_box,preprocess_mask_for_analysis,split_mask_to_contours,call_crack_api,parse_deepseek_json
#sk-8a5add1a6785414a9ff1b2653e760880
# ===== DeepSeek 初始化 =====
client = OpenAI(
    api_key="sk-8a5add1a6785414a9ff1b2653e760880",
    base_url="https://api.deepseek.com"
)
# ===== U-Net API 配置 =====
UNET_API_URL = "http://127.0.0.1:5000/unet/{model_id}/segment"

MODEL_MAPPING = {
    "unet_Fracture": {"color": (255, 0, 0)},          # 红色
    "unet_Induced_Fracture": {"color": (0, 0, 255)},  # 蓝色
    "unet_Vug": {"color": (0, 255, 0)}               # 绿色
}
# ===== 颜色映射 =====

# SAM2 专用颜色映射
SAM2_MAPPING = {
    "sam2_fracture": {"color": (0, 255, 0)},   # 紫色
    "sam2_vug": {"color": (255, 165, 0)}         # 橙色
}

def deepseek_filter_curves_safe(
    curves_metrics,
    x_points_list,
    image_height_px,
    image_width_px,
    min_points=200,
    max_retries=3,
    log_fn=None,
    strategy="default"
):
    """
    本地 Ollama 版本
    Consistency Adjudication
    """

    if not curves_metrics or not x_points_list:
        return [], [], []

    # ---------- 几何预筛选 ----------
    pre_filtered = [
        (i, m) for i, m in enumerate(curves_metrics)
        if i < len(x_points_list) and len(x_points_list[i]) >= min_points
    ]

    if not pre_filtered:
        if log_fn:
            log_fn("⚠️ 预处理结果为空")
        return [], [], []

    metrics_simple = [
        {"A": m.get("A"), "B": m.get("B")}
        for _, m in pre_filtered
    ]

    params_json = json.dumps(metrics_simple, ensure_ascii=False)

    prompt = generate_prompt(
        "filter_curves",
        curves_params=params_json,
        image_width_px=image_width_px,
        image_height_px=image_height_px,
        strategy=strategy
    )

    # ---------- LLM 重试 ----------
    for attempt in range(1, max_retries + 1):

        try:
            if log_fn:
                log_fn(f"📡 本地 LLM 曲线复核 (尝试 {attempt})")

            raw = call_local_llm(prompt, temperature=0)
            cleaned = safe_json_extract(raw)

            decision = json.loads(cleaned)

            valid_idx = decision.get("valid_curves", [])
            analysis_log = decision.get("analysis_log", [])

            curves_filtered = [
                pre_filtered[i][1]
                for i in valid_idx
                if i < len(pre_filtered)
            ]

            return curves_filtered, analysis_log, pre_filtered

        except Exception as e:
            if log_fn:
                log_fn(f"⚠️ LLM 解析失败: {e}")

    # ---------- 几何兜底 ----------
    if log_fn:
        log_fn("⚙️ 进入几何 fallback")

    full_log = []
    curves_filtered = []

    for idx, m in pre_filtered:
        A = abs(m.get("A", 0))
        B = m.get("B", 1.0)

        valid_mark = "valid" if B <= 0.05 and A <= image_height_px / 3 else "invalid"

        if valid_mark == "valid":
            curves_filtered.append(m)

        full_log.append(f"曲线{idx}: B={B}, |A|={A} => {valid_mark}")

    return curves_filtered, full_log, pre_filtered
# ===== DeepSeek 决策 =====
import json
import re
from local_llm import call_local_llm
def safe_json_extract(text):
    """
    处理 ```json 包裹 或 多余文本
    """
    text = text.strip()

    # 去掉 ```json ``` 包裹
    if text.startswith("```"):
        text = re.sub(r"```.*?\n", "", text)
        text = text.replace("```", "")

    # 提取第一个 JSON 对象
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        return match.group(0)

    return text


def deepseek_decide_models(user_input, yolo_results, strategy="default"):
    """
    本地 Ollama 版本
    Semantic-Constrained Model Selection
    """

    prompt = generate_prompt(
        "decide_models",
        user_input=user_input,
        yolo_results=yolo_results,
        strategy=strategy
    )

    try:
        raw = call_local_llm(prompt, temperature=0.2)
        cleaned = safe_json_extract(raw)

        decision = json.loads(cleaned)

        allowed = [
            "unet_Fracture",
            "unet_Induced_Fracture",
            "unet_Vug"
        ]

        models = [m for m in decision.get("models", []) if m in allowed]

        # Conservative fallback
        if not models:
            models = ["unet_Fracture"]

        parameters = decision.get("parameters", {})

        return models, parameters

    except Exception as e:
        print(f"⚠️ 本地 LLM 决策失败: {e}")
        return ["unet_Fracture"], {}
def generate_prompt(task_type, user_input=None, yolo_results=None, curves_params=None,
                    image_width_px=None, image_height_px=None, strategy="default"):
    """
    生成不同策略的 Prompt
    task_type: "decide_models" 或 "filter_curves"
    strategy: "default", "zero-shot", "few-shot", "hard-constraint"
    """
    if task_type == "decide_models":
        base = "你是一位图像分割智能助手，根据用户输入和YOLO检测结果自动选择最合适的地质分割模型。请严格返回 JSON。"
        if strategy == "default":
            prompt = f"{base}\n用户输入: {user_input}\nYOLO结果: {yolo_results}"
        elif strategy == "zero-shot":
            prompt = f"{base}\n用户输入: {user_input}\nYOLO结果: {yolo_results}\n只返回JSON，不要解释。"
        elif strategy == "few-shot":
            prompt = f"""{base}
示例：
用户输入: 裂缝
YOLO结果: {{}}
输出: {{
  "models": ["unet_Fracture"],
  "parameters": {{}}
}}
用户输入: 孔洞
YOLO结果: {{}}
输出: {{
  "models": ["unet_Vug"],
  "parameters": {{}}
}}
用户输入: 诱导缝
YOLO结果: {{}}
输出: {{
  "models": ["unet_Induced_Fracture"],
  "parameters": {{}}
}}
用户输入: 检测图片
YOLO结果: ["Fracture", "Vug"]
输出: {{"models": ["unet_Fracture", "unet_Vug"], "parameters": {{}} }}

现在用户输入: {user_input}
YOLO结果: {yolo_results}
请根据示例选择最合适的模型。
"""
        elif strategy == "hard-constraint":
            prompt = f"""你是一位图像分割智能助手，负责根据用户的自然语言意图与YOLO检测结果自动选择最合适的地质分割模型。
你的任务：
1. 如果用户输入中出现“裂缝”，仅使用 ["unet_Fracture"]；
2. 如果出现“孔洞”，仅使用 ["unet_Vug"]；
3. 如果出现“诱导缝”或“钻井诱导裂缝”，仅使用 ["unet_Induced_Fracture"]；
4. 如果用户输入为“检测图片”或“分析图片”，则根据 YOLO 检测结果自动匹配：
   - YOLO检测结果中包含 Fracture → 使用 ["unet_Fracture"]
   - YOLO检测结果中包含 Induced_Fracture → 使用 ["unet_Induced_Fracture"]
   - YOLO检测结果中包含 Vug → 使用 ["unet_Vug"]
5. 如果检测结果包含多种类型，可同时输出多个模型；
6. 如果均不符合条件，则返回 ["unet_Fracture"]。
输出格式固定为 JSON：
{{
  "models": ["unet_Fracture", "unet_Vug"],
  "parameters": {{}}
}}
请只输出JSON内容，不要解释。"""
        else:
            raise ValueError(f"未知策略: {strategy}")

    elif task_type == "filter_curves":
        base = "你是地质图像分析助手，根据曲线参数分析每条曲线是否有效。请严格输出 JSON,不要解释、不要添加说明、不要输出多余文字，只输出 JSON 对象。"
        if strategy in ["default", "zero-shot"]:
            prompt = (
                f"{base}\n"
                "你是一名地质曲线分析助手，请基于地质知识和电成像有关的知识判断每条曲线的有效性。\n\n"
                "【输出要求】\n"
                "1. 输出必须是合法的 JSON 对象，不能包含任何解释性文字、注释或额外说明。\n"
                "2. JSON 对象需包含以下三个键：\n"
                "   - \"analysis_log\": 包含每条曲线的分析说明。\n"
                "   - \"valid_curves\": 有效曲线索引列表。\n"
                "   - \"invalid_curves\": 无效曲线索引列表。\n"
                "3. 若无法判断，请返回：\n"
                "{{\"analysis_log\": [], \"valid_curves\": [], \"invalid_curves\": []}}\n\n"
                f"输入数据如下：\n图像宽度={image_width_px}, 高度={image_height_px}, 参数={curves_params}\n\n"
                "请直接输出符合要求的 JSON 对象。"
            )
        elif strategy == "few-shot":
            prompt = (
                f"你是一名地质测井分析专家，请学习以下示例的输入输出格式。\n"
                "任务：判断每条曲线是否有效，并输出 JSON 结果。\n\n"
                "【重要要求】\n"
                "- 你必须只输出 JSON 对象，不能输出任何解释、前缀、注释或自然语言。\n"
                "- 不要输出“输出:”或“结果如下:”，不要添加换行说明。\n"
                "- 输出必须严格遵守 JSON 语法。\n\n"
                "示例：\n"
                "输入: 图像宽度=472, 高度=1422, 参数=[{{\"A\":50,\"B\":0.02}}, {{\"A\":700,\"B\":0.1}},{{\"A\":-66.71,\"B\":0.0132}}, {{\"A\":44,\"B\":0.0132}}, {{\"A\":500,\"B\":0.060}}, {{\"A\":-43.17,\"B\":0.0221}}, {{\"A\":68199.68,\"B\":0.0003 }},{{\"A\":7.62,\"B\":0.0647}} ]\n"
                "输出:\n"
                "{{\n"
                "  \"analysis_log\": [\n"
                "    \"曲线0: A=50, B=0.02 ✅\",\n"
                "    \"曲线1: A=700, B=0.1 ❌\",\n"
                "    \"曲线2: B=0.0132，，A=-66.71, |A|=66.71 ✅\",\n"
                "    \"曲线3: B=0.0132 , A=44, |A|=44  ✅\",\n"
                "    \"曲线4: B=0.060, A=500, |A|=500  ❌ \",\n"
                "    \"曲线5: B=0.0221 , A=-43.17, |A|=43.17 ✅ \",\n"
                "    \"曲线6: B=0.0003 , A=68199.68, |A|=68199.68 ❌ \",\n"
                "    \"曲线7: B=0.0647, A=7.62, |A|=7.62 ❌ \"\n"
                "  ],\n"
                "  \"valid_curves\": [0, 2, 3, 5],\n"
                "  \"invalid_curves\": [1, 4, 6, 7]\n"
                "}}\n\n"
                f"现在输入:\n图像宽度={image_width_px}, 高度={image_height_px}, 参数={curves_params}\n\n"
                "请模仿示例输出，仅返回 JSON 对象。\n"
                "如果无法判断，请返回：\n"
                "{{\"analysis_log\": [], \"valid_curves\": [], \"invalid_curves\": []}}"
            )

        elif strategy == "hard-constraint":
            prompt = f"""
你是地质图像分析助手，负责根据曲线参数判断哪些曲线有效。请严格遵守以下规则生成输出：

【任务要求】
1. 对输入的每条曲线都生成一条日志（即使是无效曲线，也必须生成）。
2. 在处理振幅参数 A 时，请取其绝对值，因为振幅应为正数。
3. 判断标准：
   - 若 B <= 0.02 且 |A| <= 图像高度 * 2/3，则该曲线有效（✅）；
   - 否则视为无效（❌）。

【输出格式】
必须严格返回 JSON 格式，不允许出现除 JSON 外的任何解释或文字。示例如下：
{{
  "analysis_log": ["曲线0: A=100 ✅", "曲线1: A=600 ❌"],
  "valid_curves": [0],
  "invalid_curves": [1]
}}

【输入数据】
图像宽度 = {image_width_px}, 高度 = {image_height_px}, 参数 = {curves_params}

请严格按照上述标准输出 JSON。
"""
        else:
            raise ValueError(f"未知策略: {strategy}")
    else:
        raise ValueError(f"未知任务类型: {task_type}")

    return prompt

import os
import requests
from typing import Dict, Any, Optional


def generate_comprehensive_report(
    result: Any,
    image_path: str,
    log=print,
    report_api_url: str = "http://127.0.0.1:8095/generate_comprehensive_report",
    timeout: int = 120
) -> Dict[str, Optional[str]]:
    """
    Generate comprehensive geological report via report service.

    Args:
        result: DeepSeek raw result (string or dict)
        image_path: Path to image used for verification
        log: Logging function
        report_api_url: Report generation API endpoint
        timeout: HTTP timeout (seconds)

    Returns:
        dict:
            {
                "report_path": str or None,
                "report_preview": str or None,
                "status": "success" | "fallback" | "failed"
            }
    """
    try:
        # ---------- 1. Parse DeepSeek JSON ----------
        deepseek_result_json = parse_deepseek_json(result)

        payload = {
            "result": deepseek_result_json,
            "image_path": image_path,
            "image_url": image_path
        }

        log("🔹 调用报告生成接口 ...")
        report_resp = requests.post(
            report_api_url,
            json=payload,
            timeout=timeout
        )

        # ---------- 2. HTTP failure ----------
        if report_resp.status_code != 200:
            log(f"❌ 报告生成失败: HTTP {report_resp.status_code}")
            return {
                "report_path": None,
                "report_preview": None,
                "status": "failed"
            }

        # ---------- 3. Try JSON response ----------
        try:
            resp_json = report_resp.json()
            report_path = resp_json.get("report_path", "")
            report_preview = resp_json.get("report_preview", "")

            log("✅ 报告生成完成")
            log(f"📄 报告路径: {report_path}")
            log(f"📋 报告预览: {report_preview[:500]}...")

            return {
                "report_path": report_path,
                "report_preview": report_preview,
                "status": "success"
            }

        except Exception:
            # ---------- 4. Fallback: Markdown → Word ----------
            log("⚠️ 报告接口返回非 JSON，尝试写入 Word ...")

            report_text = report_resp.text.strip()

            from docx import Document
            from docx.shared import Inches

            doc = Document()
            doc.add_heading("电成像综合地质分析报告", level=1)

            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("#"):
                    doc.add_heading(
                        line.replace("#", "").strip(),
                        level=min(line.count("#"), 4)
                    )
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

            if image_path and os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(5.5))

            report_path = "DeepSeek_Report_TCA3.docx"
            doc.save(report_path)

            log(f"✅ 报告已写入 Word 文件: {report_path}")

            return {
                "report_path": report_path,
                "report_preview": report_text[:500],
                "status": "fallback"
            }

    except Exception as e:
        log(f"⚠️ 报告生成时出现异常: {e}")
        return {
            "report_path": None,
            "report_preview": None,
            "status": "failed"
        }


